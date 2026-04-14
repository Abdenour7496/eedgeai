"""
document_intel.py — OCR and Document Intelligence for eedgeai proxy.

Pipeline
--------
1.  Detect document category: text-PDF, scanned-PDF, image, DOCX, plain-text
2.  Extract raw text:
      • Text PDFs  → pdfminer.six  (layout-aware; preserves columns / tables)
      • Scanned PDFs → pdf2image → Tesseract OCR (per page)
      • Images     → Tesseract OCR  (already handled upstream via vision, but
                     docint can re-run for structured output)
      • DOCX       → python-docx paragraph + table extraction
3.  Structured extraction (vision model, page-image based — opt-in):
      • Table extraction     → Markdown tables
      • Form field detection → key-value pairs
4.  Text-model extraction (chat completion, cheap):
      • Document classification (invoice, report, medical, legal …)
      • Named entity extraction (dates, amounts, names, orgs)
5.  Return DocIntelResult — consumed by /api/docint and /api/ingest

Env vars
--------
  DOCINT_VISION         1/0  — enable vision-based table+form extraction (default 1)
  DOCINT_MAX_PAGES      max pages to OCR/render for vision (default 10)
  VISION_MAX_PX         max pixel side for rendered pages (default 1024)
  OPENAI_API_KEY        needed for chat + vision
  ANTHROPIC_API_KEY     used when LLM_BACKEND=anthropic
  LLM_BACKEND           openai|anthropic (default openai)
"""

from __future__ import annotations

import io
import json
import logging
import os
import tempfile
from dataclasses import dataclass, field, asdict

import httpx

log = logging.getLogger(__name__)

# ── Config ─────────────────────────────────────────────────────────────────────

DOCINT_VISION   = os.getenv("DOCINT_VISION", "1") == "1"
DOCINT_MAX_PGS  = int(os.getenv("DOCINT_MAX_PAGES", "10"))
VISION_MAX_PX   = int(os.getenv("VISION_MAX_PX", "1024"))
OPENAI_KEY      = os.getenv("OPENAI_API_KEY", "")
ANTHROPIC_KEY   = os.getenv("ANTHROPIC_API_KEY", "")
LLM_BACKEND     = os.getenv("LLM_BACKEND", "openai")
VISION_BACKEND  = os.getenv("VISION_BACKEND") or LLM_BACKEND
CHAT_MODEL      = os.getenv("OPENAI_CHAT_MODEL", "gpt-4o")
ANTHROPIC_MODEL = os.getenv("ANTHROPIC_CHAT_MODEL", "claude-sonnet-4-6")
VISION_MODEL    = os.getenv("VISION_MODEL",
                            ANTHROPIC_MODEL if VISION_BACKEND == "anthropic" else "gpt-4o")
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://ollama:11434")
OLLAMA_MODEL    = os.getenv("OLLAMA_MODEL", "llama3.2")
MAX_TOKENS_TEXT = 2048
MAX_TOKENS_VIS  = 1500

# ── Data model ─────────────────────────────────────────────────────────────────

@dataclass
class DocTable:
    caption:   str
    markdown:  str
    row_count: int
    col_count: int
    page:      int = 0

@dataclass
class DocSection:
    title:   str
    content: str
    level:   int   # 0 = body paragraph, 1-6 = heading level
    page:    int = 0

@dataclass
class DocEntities:
    dates:         list[str] = field(default_factory=list)
    amounts:       list[str] = field(default_factory=list)
    names:         list[str] = field(default_factory=list)
    organizations: list[str] = field(default_factory=list)
    identifiers:   list[str] = field(default_factory=list)
    locations:     list[str] = field(default_factory=list)

@dataclass
class DocIntelResult:
    # Core
    text:             str
    is_scanned:       bool
    pages:            int
    language:         str
    char_count:       int

    # Classification
    document_type:    str   # invoice | report | medical | form | legal | scientific | other
    document_subtype: str
    type_confidence:  float

    # Structured extraction
    tables:           list[DocTable]
    form_fields:      dict[str, str]
    entities:         DocEntities
    sections:         list[DocSection]

    # Neo4j props (flat, string values)
    image_props:      dict[str, str]

    def to_rich_text(self) -> str:
        """Combine all extracted content into a single string for chunking/embedding."""
        parts = []

        # Header block
        parts.append(
            f"[DOCUMENT: {self.document_type.upper()}"
            + (f" / {self.document_subtype}" if self.document_subtype else "")
            + f"  pages={self.pages}  lang={self.language}"
            + (" (scanned/OCR)" if self.is_scanned else "")
            + "]"
        )
        parts.append("")

        # Classification context
        if self.document_type != "other":
            parts.append(f"Document Type : {self.document_type}")
            if self.document_subtype:
                parts.append(f"Subtype       : {self.document_subtype}")
            parts.append("")

        # Entities summary
        ents = self.entities
        ent_lines = []
        if ents.dates:         ent_lines.append(f"Dates         : {', '.join(ents.dates[:10])}")
        if ents.amounts:       ent_lines.append(f"Amounts       : {', '.join(ents.amounts[:10])}")
        if ents.names:         ent_lines.append(f"Names         : {', '.join(ents.names[:10])}")
        if ents.organizations: ent_lines.append(f"Organizations : {', '.join(ents.organizations[:10])}")
        if ents.identifiers:   ent_lines.append(f"Identifiers   : {', '.join(ents.identifiers[:10])}")
        if ents.locations:     ent_lines.append(f"Locations     : {', '.join(ents.locations[:10])}")
        if ent_lines:
            parts.append("=== ENTITIES ===")
            parts.extend(ent_lines)
            parts.append("")

        # Form fields
        if self.form_fields:
            parts.append("=== FORM FIELDS ===")
            for k, v in self.form_fields.items():
                parts.append(f"{k}: {v}")
            parts.append("")

        # Tables (as Markdown — embeddings understand Markdown tables)
        if self.tables:
            parts.append("=== TABLES ===")
            for t in self.tables:
                if t.caption:
                    parts.append(f"Table: {t.caption}")
                parts.append(t.markdown)
                parts.append("")

        # Main text body
        parts.append("=== CONTENT ===")
        parts.append(self.text)

        return "\n".join(parts)

# ── LLM helpers ────────────────────────────────────────────────────────────────

async def _chat(messages: list[dict], max_tokens: int = MAX_TOKENS_TEXT) -> str:
    """Call the configured LLM with text-only messages."""
    if LLM_BACKEND == "ollama":
        async with httpx.AsyncClient(timeout=120) as c:
            r = await c.post(
                f"{OLLAMA_BASE_URL}/v1/chat/completions",
                json={"model": OLLAMA_MODEL, "max_tokens": max_tokens, "messages": messages},
            )
            r.raise_for_status()
            return r.json()["choices"][0]["message"]["content"].strip()

    if LLM_BACKEND == "anthropic" and ANTHROPIC_KEY:
        async with httpx.AsyncClient(timeout=60) as c:
            r = await c.post(
                "https://api.anthropic.com/v1/messages",
                headers={"x-api-key": ANTHROPIC_KEY,
                         "anthropic-version": "2023-06-01",
                         "Content-Type": "application/json"},
                json={"model": ANTHROPIC_MODEL, "max_tokens": max_tokens,
                      "messages": messages},
            )
            r.raise_for_status()
            return r.json()["content"][0]["text"].strip()

    if not OPENAI_KEY:
        return ""
    async with httpx.AsyncClient(timeout=60) as c:
        r = await c.post(
            "https://api.openai.com/v1/chat/completions",
            headers={"Authorization": f"Bearer {OPENAI_KEY}",
                     "Content-Type": "application/json"},
            json={"model": CHAT_MODEL, "max_tokens": max_tokens, "messages": messages},
        )
        r.raise_for_status()
        return r.json()["choices"][0]["message"]["content"].strip()


async def _vision(png_b64: str, prompt: str, max_tokens: int = MAX_TOKENS_VIS) -> str:
    """Call the vision model with a PNG image (base64) and a text prompt."""
    _vb = VISION_BACKEND   # respects VISION_BACKEND env, falls back to LLM_BACKEND

    if _vb == "ollama":
        # Use Ollama vision-capable model (e.g. llava, llama3.2-vision)
        async with httpx.AsyncClient(timeout=120) as c:
            r = await c.post(
                f"{OLLAMA_BASE_URL}/v1/chat/completions",
                json={"model": OLLAMA_MODEL, "max_tokens": max_tokens,
                      "messages": [{"role": "user", "content": [
                          {"type": "image_url", "image_url": {
                              "url": f"data:image/png;base64,{png_b64}"}},
                          {"type": "text", "text": prompt},
                      ]}]},
            )
            r.raise_for_status()
            return r.json()["choices"][0]["message"]["content"].strip()

    if _vb == "anthropic" and ANTHROPIC_KEY:
        async with httpx.AsyncClient(timeout=60) as c:
            r = await c.post(
                "https://api.anthropic.com/v1/messages",
                headers={"x-api-key": ANTHROPIC_KEY,
                         "anthropic-version": "2023-06-01",
                         "Content-Type": "application/json"},
                json={"model": VISION_MODEL, "max_tokens": max_tokens,
                      "messages": [{"role": "user", "content": [
                          {"type": "image", "source": {"type": "base64",
                           "media_type": "image/png", "data": png_b64}},
                          {"type": "text", "text": prompt},
                      ]}]},
            )
            r.raise_for_status()
            return r.json()["content"][0]["text"].strip()

    if not OPENAI_KEY:
        return ""
    async with httpx.AsyncClient(timeout=60) as c:
        r = await c.post(
            "https://api.openai.com/v1/chat/completions",
            headers={"Authorization": f"Bearer {OPENAI_KEY}",
                     "Content-Type": "application/json"},
            json={"model": VISION_MODEL, "max_tokens": max_tokens,
                  "messages": [{"role": "user", "content": [
                      {"type": "image_url", "image_url": {
                          "url": f"data:image/png;base64,{png_b64}", "detail": "high"}},
                      {"type": "text", "text": prompt},
                  ]}]},
        )
        r.raise_for_status()
        return r.json()["choices"][0]["message"]["content"].strip()


def _parse_json_response(raw: str) -> dict:
    """Extract JSON from an LLM response that may wrap it in markdown fences."""
    raw = raw.strip()
    # Strip ```json ... ``` fences
    if raw.startswith("```"):
        lines = raw.split("\n")
        raw = "\n".join(lines[1:-1] if lines[-1].strip() == "```" else lines[1:])
    try:
        return json.loads(raw)
    except Exception:
        return {}

# ── PDF helpers ────────────────────────────────────────────────────────────────

def _is_scanned_pdf(data: bytes) -> bool:
    """Return True if the PDF has negligible extractable text (image-based / scanned)."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        if not reader.pages:
            return False
        sample_pages = reader.pages[:min(5, len(reader.pages))]
        total_chars = sum(len((p.extract_text() or "").strip()) for p in sample_pages)
        avg = total_chars / len(sample_pages)
        return avg < 30   # fewer than 30 chars/page → treat as scanned
    except Exception:
        return False


def _extract_pdf_text_pdfminer(data: bytes) -> str:
    """Layout-aware PDF text extraction using pdfminer.six."""
    try:
        from pdfminer.high_level import extract_text
        from pdfminer.layout import LAParams
        return extract_text(
            io.BytesIO(data),
            laparams=LAParams(
                line_overlap=0.5,
                char_margin=2.0,
                word_margin=0.1,
                boxes_flow=0.5,
                detect_vertical=False,
                all_texts=False,
            ),
        )
    except ImportError:
        # Fall back to pypdf
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        return "\n".join(p.extract_text() or "" for p in reader.pages)


def _pdf_page_count(data: bytes) -> int:
    try:
        from pypdf import PdfReader
        return len(PdfReader(io.BytesIO(data)).pages)
    except Exception:
        return 1


def _pdf_to_pil_images(data: bytes, max_pages: int = DOCINT_MAX_PGS):
    """Convert PDF pages → PIL Images using pdf2image (requires poppler)."""
    from pdf2image import convert_from_bytes
    return convert_from_bytes(
        data, dpi=200,
        first_page=1, last_page=max_pages,
        fmt="RGB",
    )

# ── OCR ────────────────────────────────────────────────────────────────────────

def _ocr_pil_image(img, lang: str = "eng") -> str:
    """Run Tesseract OCR on a PIL Image."""
    import pytesseract
    cfg = "--psm 6 --oem 3"   # assume uniform block of text; use LSTM engine
    return pytesseract.image_to_string(img, lang=lang, config=cfg)


def _ocr_pdf_scanned(data: bytes, max_pages: int = DOCINT_MAX_PGS) -> tuple[str, int]:
    """Convert scanned PDF pages → OCR text. Returns (text, page_count)."""
    images = _pdf_to_pil_images(data, max_pages)
    pages_text = []
    for i, img in enumerate(images):
        page_text = _ocr_pil_image(img)
        pages_text.append(f"--- Page {i + 1} ---\n{page_text}")
    return "\n\n".join(pages_text), len(images)


def _pil_to_png_b64(img) -> str:
    import base64
    img.thumbnail((VISION_MAX_PX, VISION_MAX_PX))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return base64.b64encode(buf.getvalue()).decode()

# ── DOCX helpers ───────────────────────────────────────────────────────────────

def _extract_docx(data: bytes) -> tuple[str, list[DocTable], list[DocSection]]:
    """Extract text, tables, and sections from a DOCX file."""
    from docx import Document as DocxDoc
    from docx.oxml.ns import qn

    doc = DocxDoc(io.BytesIO(data))
    sections: list[DocSection] = []
    tables: list[DocTable] = []

    for block in doc.element.body:
        tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag

        if tag == "p":
            from docx.text.paragraph import Paragraph
            p = Paragraph(block, doc)
            text = p.text.strip()
            if not text:
                continue
            style = p.style.name if p.style else ""
            level = 0
            if style.startswith("Heading"):
                try:
                    level = int(style.split()[-1])
                except ValueError:
                    level = 1
            sections.append(DocSection(title=text if level else "", content="" if level else text, level=level))

        elif tag == "tbl":
            from docx.table import Table
            tbl = Table(block, doc)
            rows = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]
            if not rows:
                continue
            header = rows[0]
            sep = ["---"] * len(header)
            md_rows = [
                "| " + " | ".join(header) + " |",
                "| " + " | ".join(sep) + " |",
            ] + ["| " + " | ".join(r) + " |" for r in rows[1:]]
            tables.append(DocTable(
                caption="", markdown="\n".join(md_rows),
                row_count=len(rows) - 1, col_count=len(header)
            ))

    # Flatten sections into a single text body
    text_parts = []
    for s in sections:
        if s.level > 0 and s.title:
            text_parts.append(f"{'#' * s.level} {s.title}")
        elif s.content:
            text_parts.append(s.content)

    return "\n\n".join(text_parts), tables, sections

# ── Structured extraction via vision ──────────────────────────────────────────

async def _extract_tables_vision(pil_images: list) -> list[DocTable]:
    """Run vision model on page images to extract tables as Markdown."""
    tables: list[DocTable] = []

    for page_num, img in enumerate(pil_images[:DOCINT_MAX_PGS]):
        b64 = _pil_to_png_b64(img)
        prompt = (
            "Extract ALL tables from this document page.\n"
            "For each table return:\n"
            "  - caption: the table title or empty string\n"
            "  - markdown: the full table in Markdown pipe format\n"
            "Return ONLY valid JSON: "
            '{"tables": [{"caption": "...", "markdown": "|col1|col2|\\n|---|---|\\n|a|b|"}]}\n'
            "If no tables are present return: {\"tables\": []}"
        )
        raw = await _vision(b64, prompt)
        parsed = _parse_json_response(raw)
        for t in parsed.get("tables", []):
            md = t.get("markdown", "").strip()
            if not md:
                continue
            lines = [l for l in md.split("\n") if l.strip()]
            rows = max(0, len(lines) - 2)   # subtract header + separator
            cols = md.count("|") // max(len(lines), 1)
            tables.append(DocTable(
                caption=t.get("caption", ""),
                markdown=md,
                row_count=rows,
                col_count=cols,
                page=page_num + 1,
            ))

    return tables


async def _extract_forms_vision(pil_images: list) -> dict[str, str]:
    """Extract form fields and their values from page images."""
    fields: dict[str, str] = {}

    for img in pil_images[:min(3, DOCINT_MAX_PGS)]:
        b64 = _pil_to_png_b64(img)
        prompt = (
            "Extract all form fields and their filled values from this document page.\n"
            "Return ONLY valid JSON: {\"fields\": {\"Field Name\": \"value\"}}\n"
            "Examples: {\"Invoice Number\": \"INV-001\", \"Date\": \"2024-01-15\"}\n"
            "If no form fields exist, return: {\"fields\": {}}"
        )
        raw = await _vision(b64, prompt)
        parsed = _parse_json_response(raw)
        fields.update(parsed.get("fields", {}))

    return fields

# ── Text-model extraction (no vision cost) ────────────────────────────────────

async def _classify_document(text_preview: str) -> tuple[str, str, float]:
    """Classify document type using LLM on extracted text. Returns (type, subtype, confidence)."""
    prompt = (
        "Classify the following document. Return ONLY valid JSON:\n"
        '{"type": "...", "subtype": "...", "confidence": 0.95, "language": "en"}\n\n'
        "type must be one of: invoice, receipt, medical_record, lab_report, imaging_report, "
        "legal, contract, scientific, news, form, report, email, memo, technical, other\n"
        "subtype is a more specific description (e.g. 'blood_test', 'chest_ct', 'purchase_order')\n\n"
        f"Document text (first 1500 chars):\n{text_preview[:1500]}"
    )
    raw = await _chat([{"role": "user", "content": prompt}])
    parsed = _parse_json_response(raw)
    return (
        parsed.get("type", "other"),
        parsed.get("subtype", ""),
        float(parsed.get("confidence", 0.5)),
    )


async def _extract_entities(text: str) -> DocEntities:
    """Extract named entities from text using the LLM."""
    prompt = (
        "Extract named entities from the following text. Return ONLY valid JSON:\n"
        '{"dates": [], "amounts": [], "names": [], "organizations": [], '
        '"identifiers": [], "locations": []}\n\n'
        "- dates: all dates and date ranges\n"
        "- amounts: monetary values, measurements, quantities\n"
        "- names: person names\n"
        "- organizations: company, hospital, institution names\n"
        "- identifiers: IDs, reference numbers, serial numbers\n"
        "- locations: addresses, cities, countries\n\n"
        f"Text:\n{text[:3000]}"
    )
    raw = await _chat([{"role": "user", "content": prompt}])
    parsed = _parse_json_response(raw)
    return DocEntities(
        dates=parsed.get("dates", []),
        amounts=parsed.get("amounts", []),
        names=parsed.get("names", []),
        organizations=parsed.get("organizations", []),
        identifiers=parsed.get("identifiers", []),
        locations=parsed.get("locations", []),
    )


async def _detect_language(text_sample: str) -> str:
    """Detect language from text sample via LLM."""
    if not text_sample.strip():
        return "unknown"
    prompt = (
        f'What language is this text written in? Reply with only the ISO 639-1 code (e.g. "en", "fr", "de").\n'
        f"Text: {text_sample[:300]}"
    )
    raw = await _chat([{"role": "user", "content": prompt}])
    code = raw.strip().lower().replace('"', "").replace("'", "")[:5]
    return code if code.isalpha() else "en"

# ── Main entry point ───────────────────────────────────────────────────────────

async def process_document(
    filename: str,
    data: bytes,
    *,
    extract_tables: bool = True,
    extract_forms: bool = True,
    extract_entities: bool = True,
    classify: bool = True,
    vision: bool | None = None,   # None = use DOCINT_VISION env default
) -> DocIntelResult:
    """
    Run the full Document Intelligence pipeline on a file.

    Returns a DocIntelResult ready for:
      - result.to_rich_text()  → ingest into Qdrant/Neo4j
      - result.tables          → structured table data
      - result.form_fields     → key-value pairs
      - result.entities        → named entities
      - result.image_props     → Neo4j Document node properties
    """
    do_vision = DOCINT_VISION if vision is None else vision
    ext = os.path.splitext(filename)[1].lower()
    if filename.lower().endswith(".nii.gz"):
        ext = ".nii.gz"

    text = ""
    is_scanned = False
    pages = 1
    tables: list[DocTable] = []
    form_fields: dict[str, str] = {}
    pil_images: list = []

    # ── Step 1: Text extraction ────────────────────────────────────────────────

    if ext == ".pdf":
        pages = _pdf_page_count(data)
        is_scanned = _is_scanned_pdf(data)

        if is_scanned:
            log.info("DocInt: scanned PDF (%d pages) → OCR", pages)
            text, ocr_pages = _ocr_pdf_scanned(data, DOCINT_MAX_PGS)
            pages = max(pages, ocr_pages)
            if do_vision:
                try:
                    pil_images = _pdf_to_pil_images(data, DOCINT_MAX_PGS)
                except Exception as e:
                    log.warning("DocInt: pdf2image failed: %s", e)
        else:
            log.info("DocInt: text PDF (%d pages) → pdfminer", pages)
            text = _extract_pdf_text_pdfminer(data)
            if do_vision and (extract_tables or extract_forms):
                try:
                    pil_images = _pdf_to_pil_images(data, DOCINT_MAX_PGS)
                except Exception as e:
                    log.warning("DocInt: pdf2image failed: %s", e)

    elif ext == ".docx":
        log.info("DocInt: DOCX → python-docx")
        text, tables, _ = _extract_docx(data)

    elif ext in {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp",
                 ".tiff", ".tif", ".avif", ".dcm", ".dicom", ".nii", ".nii.gz"}:
        # For raw image/medical files: OCR for text layer
        log.info("DocInt: image/medical → Tesseract OCR")
        try:
            from PIL import Image
            img = Image.open(io.BytesIO(data))
            text = _ocr_pil_image(img)
            pil_images = [img]
            is_scanned = True
        except Exception as e:
            log.warning("DocInt: Tesseract OCR failed: %s", e)
            text = f"[OCR failed: {e}]"

    else:
        # Plain text, JSON, CSV, MD, etc.
        text = data.decode("utf-8", errors="replace")

    text = text.strip()

    # ── Step 2: Vision-based structure extraction ──────────────────────────────

    if do_vision and pil_images:
        if extract_tables and not tables:   # DOCX already extracted tables
            log.info("DocInt: extracting tables via vision (%d pages)", len(pil_images))
            try:
                tables = await _extract_tables_vision(pil_images)
            except Exception as e:
                log.warning("DocInt: table vision failed: %s", e)

        if extract_forms:
            log.info("DocInt: extracting form fields via vision")
            try:
                form_fields = await _extract_forms_vision(pil_images)
            except Exception as e:
                log.warning("DocInt: form vision failed: %s", e)

    # ── Step 3: Text-model extraction ─────────────────────────────────────────

    text_preview = text[:2000]
    doc_type, doc_subtype, type_conf = "other", "", 0.5
    entities = DocEntities()
    language = "en"

    if text_preview:
        tasks = []

        if classify:
            try:
                doc_type, doc_subtype, type_conf = await _classify_document(text_preview)
            except Exception as e:
                log.warning("DocInt: classification failed: %s", e)

        if extract_entities:
            try:
                entities = await _extract_entities(text)
            except Exception as e:
                log.warning("DocInt: entity extraction failed: %s", e)

        try:
            language = await _detect_language(text_preview)
        except Exception as e:
            log.warning("DocInt: language detection failed: %s", e)

    # ── Step 4: Build result ───────────────────────────────────────────────────

    image_props = {
        "docint_type":       doc_type,
        "docint_subtype":    doc_subtype,
        "docint_confidence": str(round(type_conf, 3)),
        "docint_is_scanned": str(is_scanned),
        "docint_pages":      str(pages),
        "docint_language":   language,
        "docint_tables":     str(len(tables)),
        "docint_form_fields": str(len(form_fields)),
        "docint_entities":   str(
            len(entities.dates) + len(entities.amounts) +
            len(entities.names) + len(entities.organizations)
        ),
    }

    result = DocIntelResult(
        text=text,
        is_scanned=is_scanned,
        pages=pages,
        language=language,
        char_count=len(text),
        document_type=doc_type,
        document_subtype=doc_subtype,
        type_confidence=type_conf,
        tables=tables,
        form_fields=form_fields,
        entities=entities,
        sections=[],
        image_props=image_props,
    )

    log.info(
        "DocInt: done — type=%s, scanned=%s, pages=%d, tables=%d, fields=%d",
        doc_type, is_scanned, pages, len(tables), len(form_fields),
    )
    return result
