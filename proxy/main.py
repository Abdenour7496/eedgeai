"""
Graph-Centric Orchestrated Retrieval (GCOR) proxy — Cognitive Infrastructure.

Architecture
────────────
  Neo4j   = cognitive backbone   (primary truth, structure, reasoning)
  Qdrant  = semantic perception  (index over Neo4j nodes, similarity only)

─────────────────────────────────────────────────────────────────────────────
Cognitive infrastructure additions over baseline GCOR:

  Confidence filtering
    • Qdrant hits below CONFIDENCE_THRESHOLD are dropped before expansion
    • Neo4j expansion uses $min_conf to filter Memory/Inference/Belief

  Temporal validity
    • Qdrant hits with expired valid_to are dropped (dead knowledge)
    • Neo4j expansion passes $now so queries respect valid_to

  Agent partitioning
    • AGENT_ID env var scopes Memory/Belief/Inference to one agent
    • Access-control filter drops hits whose access_level forbids $AGENT_ID

  New intents
    • inference   → resolve :Inference chains and their evidence
    • belief      → retrieve :Belief nodes and contradiction graph

─────────────────────────────────────────────────────────────────────────────
Retrieval flow for every /v1/chat/completions request
─────────────────────────────────────────────────────
  1. INTENT CLASSIFICATION
     factual | planning | dependency | memory | semantic | inference | belief

  2. SEMANTIC PHASE  (Qdrant)
     Embed → top-K candidate nodes
     Filter by: score, confidence, temporal validity, access control

  3. STRUCTURAL PHASE  (Neo4j)
     Intent-specific Cypher with confidence + temporal + agent_id params

  4. REFLECTION CHECK
     Neo4j empty but Qdrant has hits → fall back to chunk text
     Both empty → LLM general knowledge

  5. CONTEXT INJECTION
     Structured system message including confidence scores

Other endpoints
───────────────
  GET  /v1/models      → live OpenAI model list (with fallback)
  POST /v1/embeddings  → proxied to OpenAI
  GET  /health         → liveness check
"""

import hashlib
import io
import json
import logging
import os
import time
import uuid
from datetime import datetime, timezone

import httpx
from fastapi import BackgroundTasks, FastAPI, File, Form, HTTPException, Query, Request, UploadFile
from fastapi.responses import HTMLResponse, JSONResponse, RedirectResponse, StreamingResponse
from neo4j import AsyncGraphDatabase
from prometheus_client import Counter, Histogram
from prometheus_fastapi_instrumentator import Instrumentator
from document_intel import process_document, DocIntelResult

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)

app = FastAPI()

# ── Prometheus metrics ─────────────────────────────────────────────────────────

# Auto-instruments HTTP request count + latency for every FastAPI endpoint
Instrumentator().instrument(app).expose(app)

# GCOR RAG pipeline
METRIC_RAG_REQUESTS = Counter(
    "gcor_rag_requests_total",
    "Total GCOR RAG pipeline invocations",
    ["intent", "fallback_mode"],
)
METRIC_RAG_DURATION = Histogram(
    "gcor_rag_duration_seconds",
    "End-to-end GCOR pipeline latency (embed → Qdrant → Neo4j → context build)",
    buckets=[0.1, 0.25, 0.5, 1.0, 2.0, 5.0, 10.0, 30.0],
)
METRIC_QDRANT_HITS = Histogram(
    "gcor_qdrant_hits",
    "Number of Qdrant hits returned per RAG request (after cognitive filters)",
    buckets=[0, 1, 2, 3, 4, 5, 6, 8, 10, 15, 20],
)
METRIC_NEO4J_RECORDS = Histogram(
    "gcor_neo4j_records",
    "Number of Neo4j graph records expanded per RAG request",
    buckets=[0, 1, 2, 5, 10, 15, 20, 30, 50],
)

# Embedding
METRIC_EMBED_DURATION = Histogram(
    "gcor_embed_duration_seconds",
    "OpenAI embedding API call latency",
    buckets=[0.05, 0.1, 0.2, 0.5, 1.0, 2.0, 5.0],
)

# LLM calls
METRIC_LLM_REQUESTS = Counter(
    "gcor_llm_requests_total",
    "Total LLM API calls",
    ["backend", "status"],
)
METRIC_LLM_DURATION = Histogram(
    "gcor_llm_duration_seconds",
    "LLM API call latency (non-streaming only)",
    ["backend"],
    buckets=[0.5, 1.0, 2.0, 5.0, 10.0, 20.0, 40.0, 60.0, 120.0],
)

# Ingest
METRIC_INGEST_TOTAL = Counter(
    "gcor_ingest_total",
    "Total document ingest operations",
    ["status"],
)
METRIC_INGEST_CHUNKS = Histogram(
    "gcor_ingest_chunks",
    "Number of chunks produced per ingested document",
    buckets=[1, 5, 10, 20, 50, 100, 200, 500],
)

# Collection management
METRIC_COLLECTION_OPS = Counter(
    "gcor_collection_ops_total",
    "Qdrant collection management operations",
    ["operation"],
)

# ── LLM backends ───────────────────────────────────────────────────────────────
OPENAI_API_KEY       = os.getenv("OPENAI_API_KEY", "")
OPENAI_CHAT_MODEL    = os.getenv("OPENAI_CHAT_MODEL", "gpt-4o")
ANTHROPIC_API_KEY    = os.getenv("ANTHROPIC_API_KEY", "")
ANTHROPIC_CHAT_MODEL = os.getenv("ANTHROPIC_CHAT_MODEL", "claude-sonnet-4-6")
LLM_BACKEND          = os.getenv("LLM_BACKEND", "openai")   # openai | anthropic | ollama | openclaw

# ── OpenClaw ────────────────────────────────────────────────────────────────────
OPENCLAW_BASE_URL      = os.getenv("OPENCLAW_BASE_URL",      "http://openclaw:18799/v1")
OPENCLAW_GATEWAY_TOKEN = os.getenv("OPENCLAW_GATEWAY_TOKEN", "")

# ── GitHub Copilot (OpenClaw fallback) ──────────────────────────────────────────
COPILOT_API_KEY    = os.getenv("COPILOT_API_KEY", "")          # GitHub token with Copilot access
COPILOT_BASE_URL   = os.getenv("COPILOT_BASE_URL", "https://api.githubcopilot.com")
COPILOT_CHAT_MODEL = os.getenv("COPILOT_CHAT_MODEL", "gpt-4.1")

# ── Ollama ─────────────────────────────────────────────────────────────────────
OLLAMA_BASE_URL        = os.getenv("OLLAMA_BASE_URL",        "http://ollama:11434")
OLLAMA_MODEL           = os.getenv("OLLAMA_MODEL",           "llama3.2")
OLLAMA_EMBEDDING_MODEL = os.getenv("OLLAMA_EMBEDDING_MODEL", "nomic-embed-text")

# ── Embedding ──────────────────────────────────────────────────────────────────
EMBEDDING_BACKEND = os.getenv("EMBEDDING_BACKEND", "openai")   # openai | ollama
EMBEDDING_MODEL   = os.getenv("EMBEDDING_MODEL",   "text-embedding-3-small")
EMBED_BATCH_SIZE  = int(os.getenv("EMBED_BATCH_SIZE",  "32"))   # chunks per OpenAI call
EMBED_BATCH_DELAY = float(os.getenv("EMBED_BATCH_DELAY", "0.5")) # seconds between batches
EMBED_MAX_RETRIES = int(os.getenv("EMBED_MAX_RETRIES",  "6"))    # 429/5xx retry attempts
LLM_MAX_RETRIES   = int(os.getenv("LLM_MAX_RETRIES",    "3"))    # 429/529/5xx chat retry attempts
LLM_RETRY_BASE    = float(os.getenv("LLM_RETRY_BASE",  "1.0"))   # base delay for LLM retries

# ── Qdrant (semantic perception layer) ────────────────────────────────────────
QDRANT_URL        = os.getenv("QDRANT_URL",        "http://qdrant:6333")
QDRANT_COLLECTION = os.getenv("QDRANT_COLLECTION", "documents")
QDRANT_TOP_K      = int(os.getenv("QDRANT_TOP_K",  "8"))

# Cached embedding dimension — probed once at first use
_embed_dim_cache: int | None = None


async def _get_embed_dim() -> int:
    """Return the vector dimension of the current embedding model (cached)."""
    global _embed_dim_cache
    if _embed_dim_cache is not None:
        return _embed_dim_cache
    vec = await embed_text("dimension probe")
    _embed_dim_cache = len(vec) if vec else 768
    return _embed_dim_cache

# ── Neo4j (cognitive backbone) ────────────────────────────────────────────────
NEO4J_URI      = os.getenv("NEO4J_URI",      "bolt://neo4j:7687")
NEO4J_USER     = os.getenv("NEO4J_USER",     "neo4j")
NEO4J_PASSWORD = os.getenv("NEO4J_PASSWORD", "test1234")
NEO4J_DATABASE = os.getenv("NEO4J_DATABASE", "neo4j")

ENABLE_RAG = os.getenv("ENABLE_RAG", "true").lower() in ("true", "1", "yes")

# ── Cognitive infrastructure knobs ────────────────────────────────────────────
# Minimum confidence (0.0–1.0) — hits below this are discarded
CONFIDENCE_THRESHOLD = float(os.getenv("CONFIDENCE_THRESHOLD", "0.0"))
# Agent id scopes memory/belief/inference retrieval to a single partition
AGENT_ID = os.getenv("AGENT_ID", "")


# ── Step 1: Intent classification ─────────────────────────────────────────────

_INTENT_KEYWORDS = {
    "planning":    ["plan", "schedule", "next step", "roadmap", "workflow",
                    "milestone", "sprint", "agenda", "timeline", "sequence"],
    "dependency":  ["depends", "depend on", "requires", "blocked", "prerequisite", "constraint",
                    "relies on", "linked to", "what does", "what do"],
    "memory":      ["remember", "recall", "history", "last time", "previously",
                    "past", "earlier", "before we", "you said", "i told"],
    "factual":     ["what is", "how does", "why", "explain", "define", "describe",
                    "tell me about", "how do", "what are"],
    "inference":   ["infer", "conclude", "deduce", "reasoning", "derive",
                    "implication", "therefore", "follows that", "suggests",
                    "implies", "logical", "because of"],
    "belief":      ["believe", "think", "opinion", "doubt", "uncertain",
                    "probably", "likely", "assume", "suppose", "suspect",
                    "my view", "i think"],
}

# ── Embedding-based intent classification ─────────────────────────────────────

_INTENT_EXEMPLARS = {
    "factual":    "What is this? How does it work? Explain and define the concept.",
    "planning":   "What are the next steps? Create a roadmap and schedule milestones.",
    "dependency": "What does this depend on? What blocks progress? What are prerequisites?",
    "memory":     "What did we discuss before? Recall our previous conversation history.",
    "inference":  "What can we conclude? Derive the logical implication from this evidence.",
    "belief":     "What do you think? What is your opinion? What do you believe is likely?",
    "semantic":   "Search for relevant information about this topic.",
}

_INTENT_EXEMPLAR_VECS: dict = {}


def _cosine_sim(a: list, b: list) -> float:
    """Cosine similarity between two equal-length vectors."""
    if not a or not b or len(a) != len(b):
        return 0.0
    dot = sum(x * y for x, y in zip(a, b))
    norm_a = sum(x * x for x in a) ** 0.5
    norm_b = sum(x * x for x in b) ** 0.5
    if norm_a == 0 or norm_b == 0:
        return 0.0
    return dot / (norm_a * norm_b)


@app.on_event("startup")
async def _precompute_intent_exemplars():
    """Embed intent exemplars at startup for embedding-based classification."""
    global _INTENT_EXEMPLAR_VECS
    try:
        for intent, sentence in _INTENT_EXEMPLARS.items():
            vec = await embed_text(sentence)
            if vec:
                _INTENT_EXEMPLAR_VECS[intent] = vec
        logger.info("Intent exemplar embeddings precomputed: %d intents", len(_INTENT_EXEMPLAR_VECS))
    except Exception as exc:
        logger.warning("Could not precompute intent exemplars (will use keyword fallback): %s", exc)


async def classify_intent(query: str) -> str:
    """Classify query intent via embedding cosine sim; falls back to keyword scan."""
    # Embedding-based path
    if _INTENT_EXEMPLAR_VECS:
        try:
            query_vec = await embed_text(query)
            if query_vec:
                scores = {
                    intent: _cosine_sim(query_vec, vec)
                    for intent, vec in _INTENT_EXEMPLAR_VECS.items()
                }
                best = max(scores, key=scores.get)
                logger.debug("Intent scores: %s → %s", scores, best)
                return best
        except Exception as exc:
            logger.warning("Embedding intent classification failed, using keyword fallback: %s", exc)
    # Keyword fallback
    q = query.lower()
    for intent, keywords in _INTENT_KEYWORDS.items():
        if any(kw in q for kw in keywords):
            return intent
    return "semantic"


# ── Step 2a: Semantic phase (Qdrant) ──────────────────────────────────────────

def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


import asyncio
import random

async def _embed_with_retry(
    payload: dict,
    *,
    max_retries: int = EMBED_MAX_RETRIES,
) -> dict:
    """
    POST to OpenAI /v1/embeddings with exponential-backoff retry on 429 / 5xx.

    Respects the Retry-After header when present.
    Raises httpx.HTTPStatusError on permanent failure.
    """
    base_delay = 1.0
    for attempt in range(max_retries + 1):
        async with httpx.AsyncClient(timeout=60) as c:
            resp = await c.post(
                "https://api.openai.com/v1/embeddings",
                headers={"Authorization": f"Bearer {OPENAI_API_KEY}",
                         "Content-Type": "application/json"},
                json=payload,
            )

        if resp.status_code == 200:
            return resp.json()

        retry_after = None
        if resp.status_code == 429 or resp.status_code >= 500:
            # Honour Retry-After header if present
            ra = resp.headers.get("retry-after") or resp.headers.get("x-ratelimit-reset-requests")
            if ra:
                try:
                    retry_after = float(ra)
                except ValueError:
                    pass

            if attempt < max_retries:
                wait = retry_after if retry_after else (base_delay * (2 ** attempt) + random.uniform(0, 0.5))
                logger.warning(
                    "OpenAI embeddings %s (attempt %d/%d) — waiting %.1fs",
                    resp.status_code, attempt + 1, max_retries, wait,
                )
                await asyncio.sleep(wait)
                continue

        resp.raise_for_status()   # permanent error

    resp.raise_for_status()


async def _embed_batch(texts: list[str]) -> list[list[float]]:
    """Embed a batch of texts via the configured backend (OpenAI or Ollama)."""
    if EMBEDDING_BACKEND == "ollama":
        async with httpx.AsyncClient(timeout=120) as c:
            resp = await c.post(
                f"{OLLAMA_BASE_URL}/v1/embeddings",
                json={"input": texts, "model": OLLAMA_EMBEDDING_MODEL},
            )
            resp.raise_for_status()
        items = resp.json()["data"]
        items.sort(key=lambda x: x["index"])
        return [d["embedding"] for d in items]
    data = await _embed_with_retry({"input": texts, "model": EMBEDDING_MODEL})
    items = data["data"]
    items.sort(key=lambda x: x["index"])
    return [d["embedding"] for d in items]


async def embed_text(text: str) -> list | None:
    try:
        t0 = time.monotonic()
        if EMBEDDING_BACKEND == "ollama":
            async with httpx.AsyncClient(timeout=60) as c:
                resp = await c.post(
                    f"{OLLAMA_BASE_URL}/v1/embeddings",
                    json={"input": text[:8000], "model": OLLAMA_EMBEDDING_MODEL},
                )
                resp.raise_for_status()
            vec = resp.json()["data"][0]["embedding"]
        else:
            if not OPENAI_API_KEY:
                return None
            data = await _embed_with_retry({"input": text[:8000], "model": EMBEDDING_MODEL})
            vec = data["data"][0]["embedding"]
        METRIC_EMBED_DURATION.observe(time.monotonic() - t0)
        return vec
    except Exception as exc:
        logger.warning("Embedding failed: %s", exc)
        return None


async def _qdrant_search_collection(
    client: httpx.AsyncClient, collection: str, vector: list, limit: int,
) -> list:
    """Search a single Qdrant collection with server-side temporal filters.

    Qdrant pre-filters out expired (valid_to < now) and not-yet-valid
    (valid_from > now) points so they never consume top-K slots.
    """
    now = _now_iso()
    payload_filter = {
        "must": [
            # Exclude expired: valid_to must be null or >= now
            {"should": [
                {"is_empty": {"key": "valid_to"}},
                {"key": "valid_to", "range": {"gte": now}},
            ]},
            # Exclude not-yet-valid: valid_from must be null or <= now
            {"should": [
                {"is_empty": {"key": "valid_from"}},
                {"key": "valid_from", "range": {"lte": now}},
            ]},
        ],
    }
    try:
        resp = await client.post(
            f"{QDRANT_URL}/collections/{collection}/points/search",
            json={
                "vector": vector,
                "limit": limit,
                "with_payload": True,
                "filter": payload_filter,
            },
        )
        if resp.status_code == 404:
            return []
        resp.raise_for_status()
        hits = resp.json().get("result", [])
        for h in hits:
            h.setdefault("_collection", collection)
        return hits
    except Exception as exc:
        logger.warning("Qdrant search on '%s' failed: %s", collection, exc)
        return []


# Collections to search during RAG (in addition to QDRANT_COLLECTION).
# Only collections whose dimension matches the active embedding model are queried.
RAG_EXTRA_COLLECTIONS = [
    c.strip()
    for c in os.getenv("RAG_EXTRA_COLLECTIONS", "buddy_memory").split(",")
    if c.strip()
]


async def qdrant_search(vector: list) -> list:
    """Search primary + extra RAG collections, fuse by normalised cosine score.

    Each collection's scores are min-max normalised to [0, 1] before merging
    so that differences in corpus size or score distributions don't bias results.
    The top QDRANT_TOP_K results across all collections are returned.
    """
    embed_dim = len(vector)
    collections = [QDRANT_COLLECTION]

    # Discover which extra collections have matching dimensions
    async with httpx.AsyncClient(timeout=10) as c:
        for name in RAG_EXTRA_COLLECTIONS:
            if name == QDRANT_COLLECTION:
                continue
            try:
                info = await c.get(f"{QDRANT_URL}/collections/{name}")
                if info.status_code != 200:
                    continue
                col_dim = (info.json().get("result", {})
                           .get("config", {}).get("params", {})
                           .get("vectors", {}).get("size"))
                if col_dim == embed_dim:
                    collections.append(name)
                else:
                    logger.debug(
                        "Skipping collection '%s': dim %s != embed dim %s",
                        name, col_dim, embed_dim,
                    )
            except Exception:
                pass

        # Search all matching collections concurrently
        per_col = max(QDRANT_TOP_K, 12)  # fetch extra so normalisation is stable
        tasks = [
            _qdrant_search_collection(c, col, vector, per_col)
            for col in collections
        ]
        results = await asyncio.gather(*tasks)

    # Min-max normalise scores per collection then merge
    all_hits: list = []
    for hits in results:
        if not hits:
            continue
        scores = [h.get("score", 0) for h in hits]
        lo, hi = min(scores), max(scores)
        span = hi - lo if hi > lo else 1.0
        for h in hits:
            h["_raw_score"] = h.get("score", 0)
            h["score"] = (h["_raw_score"] - lo) / span
        all_hits.extend(hits)

    # Sort by normalised score descending, take top-K
    all_hits.sort(key=lambda h: h["score"], reverse=True)
    return all_hits[:QDRANT_TOP_K]


def _filter_hits(hits: list) -> list:
    """
    Apply cognitive filters to raw Qdrant hits:
      1. Confidence threshold — payload.confidence below minimum is discarded
      2. Temporal validity   — expired valid_to is discarded
      3. Access control      — access_level incompatible with AGENT_ID is discarded
    """
    now = _now_iso()
    filtered = []
    for h in hits:
        p = h.get("payload") or {}

        # ── confidence ────────────────────────────────────────────────────────
        node_confidence = p.get("confidence")
        if node_confidence is not None and node_confidence < CONFIDENCE_THRESHOLD:
            logger.debug("Dropping hit %s: confidence %.2f < %.2f",
                         h.get("id"), node_confidence, CONFIDENCE_THRESHOLD)
            continue

        # ── temporal validity ─────────────────────────────────────────────────
        valid_to = p.get("valid_to")
        if valid_to and valid_to < now:
            logger.debug("Dropping hit %s: expired valid_to %s", h.get("id"), valid_to)
            continue
        valid_from = p.get("valid_from")
        if valid_from and valid_from > now:
            logger.debug("Dropping hit %s: not yet valid (valid_from %s)", h.get("id"), valid_from)
            continue

        # ── access control ────────────────────────────────────────────────────
        access_level = p.get("access_level", "public")
        if access_level == "restricted":
            logger.debug("Dropping hit %s: restricted access", h.get("id"))
            continue
        if access_level.startswith("agent:") and AGENT_ID:
            owner = access_level.split(":", 1)[1]
            if owner != AGENT_ID:
                logger.debug("Dropping hit %s: owned by %s, not %s", h.get("id"), owner, AGENT_ID)
                continue

        filtered.append(h)
    return filtered


# ── Step 2b: Structural phase (Neo4j) ─────────────────────────────────────────

_EXPAND_CYPHER = {
    "factual": """
        MATCH (n) WHERE elementId(n) IN $ids
          AND (n.valid_from IS NULL OR n.valid_from <= $now)
          AND (n.valid_to   IS NULL OR n.valid_to   >= $now)
        OPTIONAL MATCH (n)-[r]-(related)
          WHERE (related.valid_to IS NULL OR related.valid_to >= $now)
        OPTIONAL MATCH (n)<-[:CONTAINS]-(doc:Document)
        RETURN
            properties(n) AS node, labels(n) AS labels,
            properties(doc) AS document,
            collect(DISTINCT {
                rel: type(r), props: properties(related), lbls: labels(related)
            })[..6] AS related
        LIMIT 20
    """,
    "planning": """
        MATCH (n) WHERE elementId(n) IN $ids
          AND (n.valid_from IS NULL OR n.valid_from <= $now)
          AND (n.valid_to   IS NULL OR n.valid_to   >= $now)
        OPTIONAL MATCH (n)-[:DEPENDS_ON*1..2]->(dep)
        OPTIONAL MATCH (n)<-[:CONTAINS]-(doc:Document)
        OPTIONAL MATCH (g:Goal)-[:DEPENDS_ON]->(n)
            WHERE (g.valid_to IS NULL OR g.valid_to >= $now)
        OPTIONAL MATCH (infer_node:Inference)-[:SUPPORTS]->(n)
            WHERE infer_node.confidence >= $min_conf
              AND (infer_node.valid_to IS NULL OR infer_node.valid_to >= $now)
        RETURN
            properties(n) AS node, labels(n) AS labels,
            properties(doc) AS document,
            collect(DISTINCT properties(dep))[..5]       AS dependencies,
            collect(DISTINCT properties(g))[..5]         AS blocking_goals,
            collect(DISTINCT properties(infer_node))[..3] AS supporting_inferences
        LIMIT 20
    """,
    "dependency": """
        MATCH (n) WHERE elementId(n) IN $ids
          AND (n.valid_from IS NULL OR n.valid_from <= $now)
          AND (n.valid_to   IS NULL OR n.valid_to   >= $now)
        OPTIONAL MATCH path = (n)-[:DEPENDS_ON*1..3]->(dep)
        OPTIONAL MATCH (n)<-[:CONTAINS]-(doc:Document)
        RETURN
            properties(n)  AS node, labels(n) AS labels,
            properties(doc) AS document,
            [x IN nodes(path) | properties(x)][..8] AS dependency_chain
        LIMIT 20
    """,
    "memory": """
        MATCH (n) WHERE elementId(n) IN $ids
          AND (n.valid_from IS NULL OR n.valid_from <= $now)
          AND (n.valid_to   IS NULL OR n.valid_to   >= $now)
        OPTIONAL MATCH (n)-[:MENTIONS]->(entity:Entity)
        OPTIONAL MATCH (mem:Memory)-[:ABOUT]->(entity)
            WHERE ($agent_id = '' OR mem.agent_id = $agent_id)
              AND mem.confidence >= $min_conf
              AND (mem.valid_from IS NULL OR mem.valid_from <= $now)
              AND (mem.valid_to   IS NULL OR mem.valid_to   >= $now)
        OPTIONAL MATCH (bel:Belief)-[:ABOUT]->(entity)
            WHERE ($agent_id = '' OR bel.agent_id = $agent_id)
              AND bel.confidence >= $min_conf
              AND (bel.valid_from IS NULL OR bel.valid_from <= $now)
              AND (bel.valid_to   IS NULL OR bel.valid_to   >= $now)
        OPTIONAL MATCH (evt:Event)-[:INVOLVES]->(entity)
        OPTIONAL MATCH (n)<-[:CONTAINS]-(doc:Document)
        RETURN
            properties(n)   AS node, labels(n) AS labels,
            properties(doc) AS document,
            collect(DISTINCT properties(entity))[..5] AS concepts,
            collect(DISTINCT properties(mem))[..5]     AS memories,
            collect(DISTINCT properties(bel))[..3]     AS beliefs,
            collect(DISTINCT properties(evt))[..5]     AS events
        ORDER BY evt.timestamp DESC
        LIMIT 20
    """,
    "semantic": """
        MATCH (n) WHERE elementId(n) IN $ids
          AND (n.valid_from IS NULL OR n.valid_from <= $now)
          AND (n.valid_to   IS NULL OR n.valid_to   >= $now)
        OPTIONAL MATCH (n)-[:MENTIONS]->(entity:Entity)
        OPTIONAL MATCH (n)-[:NEXT]-(neighbor:Chunk)
          WHERE (neighbor.valid_to IS NULL OR neighbor.valid_to >= $now)
        OPTIONAL MATCH (n)<-[:CONTAINS]-(doc:Document)
        RETURN
            properties(n)   AS node, labels(n) AS labels,
            properties(doc) AS document,
            collect(DISTINCT properties(entity))[..5]  AS concepts,
            collect(DISTINCT properties(neighbor))[..3] AS neighbors
        LIMIT 20
    """,
    "inference": """
        MATCH (n) WHERE elementId(n) IN $ids
          AND (n.valid_from IS NULL OR n.valid_from <= $now)
          AND (n.valid_to   IS NULL OR n.valid_to   >= $now)
        OPTIONAL MATCH (n)-[:DERIVED_FROM]->(src)
        OPTIONAL MATCH (n)-[:SUPPORTS]->(tgt)
        OPTIONAL MATCH (n)<-[:CONTAINS]-(doc:Document)
        OPTIONAL MATCH (downstream:Inference)-[:DERIVED_FROM]->(n)
            WHERE downstream.confidence >= $min_conf
              AND ($agent_id = '' OR downstream.agent_id = $agent_id)
              AND (downstream.valid_from IS NULL OR downstream.valid_from <= $now)
              AND (downstream.valid_to   IS NULL OR downstream.valid_to   >= $now)
        RETURN
            properties(n)   AS node, labels(n) AS labels,
            properties(doc) AS document,
            collect(DISTINCT properties(src))[..5]        AS sources,
            collect(DISTINCT properties(tgt))[..5]        AS supports,
            collect(DISTINCT properties(downstream))[..3] AS downstream_inferences
        LIMIT 20
    """,
    "belief": """
        MATCH (n) WHERE elementId(n) IN $ids
          AND (n.valid_from IS NULL OR n.valid_from <= $now)
          AND (n.valid_to   IS NULL OR n.valid_to   >= $now)
        OPTIONAL MATCH (a:Agent)-[:HOLDS]->(n)
        OPTIONAL MATCH (n)-[:ABOUT]->(entity:Entity)
        OPTIONAL MATCH (n)-[:CONTRADICTS]-(other:Belief)
            WHERE other.confidence >= $min_conf
              AND (other.valid_to IS NULL OR other.valid_to >= $now)
        OPTIONAL MATCH (infer_node:Inference)-[:SUPPORTS]->(n)
            WHERE infer_node.confidence >= $min_conf
              AND (infer_node.valid_from IS NULL OR infer_node.valid_from <= $now)
              AND (infer_node.valid_to   IS NULL OR infer_node.valid_to   >= $now)
        RETURN
            properties(n)   AS node, labels(n) AS labels, null AS document,
            properties(a)   AS agent,
            collect(DISTINCT properties(entity))[..5]       AS concepts,
            collect(DISTINCT properties(other))[..3]         AS contradictions,
            collect(DISTINCT properties(infer_node))[..3]    AS supporting_inferences
        LIMIT 20
    """,
}

_FALLBACK_CYPHER = """
    MATCH (n)
    WHERE any(k IN keys(n) WHERE NOT valueType(n[k]) STARTS WITH 'LIST' AND toLower(toString(n[k])) CONTAINS toLower($q))
      AND coalesce(n.confidence, 1.0) >= $min_conf
      AND (n.valid_from IS NULL OR n.valid_from <= $now)
      AND (n.valid_to   IS NULL OR n.valid_to   >= $now)
    RETURN properties(n) AS node, labels(n) AS labels, null AS document
    LIMIT 10
"""


async def neo4j_expand(
    element_ids: list,
    intent: str,
    query: str,
) -> list:
    """Expand the graph around Qdrant candidate nodes based on intent."""
    cypher = _EXPAND_CYPHER.get(intent, _EXPAND_CYPHER["semantic"])
    now = _now_iso()
    params = {
        "ids":      element_ids,
        "agent_id": AGENT_ID,
        "min_conf": CONFIDENCE_THRESHOLD,
        "now":      now,
    }

    if not element_ids:
        cypher = _FALLBACK_CYPHER
        params  = {"q": query[:100], "min_conf": CONFIDENCE_THRESHOLD, "now": now}

    try:
        driver = AsyncGraphDatabase.driver(NEO4J_URI, auth=(NEO4J_USER, NEO4J_PASSWORD))
        try:
            async with driver.session(database=NEO4J_DATABASE) as session:
                result = await session.run(cypher, params)
                return await result.data()
        finally:
            await driver.close()
    except Exception as exc:
        logger.warning("Neo4j expansion failed: %s", exc)
        return []


# ── Step 2c: Retrieval Provenance Chain ───────────────────────────────────────

async def build_provenance_chain(element_ids: list) -> str:
    """
    Trace the retrieval path: Chunk → Concept/Entity/Inference in Neo4j.
    Returns a human-readable provenance string, or "" if nothing found.
    """
    if not element_ids:
        return ""
    try:
        driver = AsyncGraphDatabase.driver(NEO4J_URI, auth=(NEO4J_USER, NEO4J_PASSWORD))
        try:
            async with driver.session(database=NEO4J_DATABASE) as session:
                result = await session.run(
                    """
                    MATCH path = (c:Chunk)-[:MENTIONS*0..2]->(n)
                    WHERE elementId(c) IN $ids
                    RETURN [node IN nodes(path) |
                        labels(node)[0] + ': ' +
                        coalesce(node.text, node.name, node.title, '?')
                    ] AS chain
                    LIMIT 5
                    """,
                    ids=element_ids,
                )
                rows = await result.data()
        finally:
            await driver.close()
        if not rows:
            return ""
        chains = [" → ".join(r["chain"][:80] if isinstance(r["chain"], list) else [str(r["chain"])]) for r in rows[:3]]
        return " | ".join(c for c in chains if c)
    except Exception as exc:
        logger.debug("Provenance chain query failed: %s", exc)
        return ""


# ── Step 3: Reflection ─────────────────────────────────────────────────────────

def _reflection_fallback(qdrant_hits: list, graph_records: list) -> list:
    if graph_records:
        return graph_records
    if qdrant_hits:
        logger.info("Reflection: Neo4j empty, falling back to Qdrant chunk text.")
        return [
            {
                "node": {
                    "text":        h["payload"].get("text", ""),
                    "document_id": h["payload"].get("document_id", ""),
                    "confidence":  h["payload"].get("confidence"),
                    "valid_from":  h["payload"].get("valid_from"),
                    "valid_to":    h["payload"].get("valid_to"),
                },
                "labels":   [h["payload"].get("node_type", "Chunk")],
                "document": {"title": h["payload"].get("document_title", "")},
            }
            for h in qdrant_hits
            if h.get("payload", {}).get("text")
        ]
    return []


# ── Context builder ────────────────────────────────────────────────────────────

def _confidence_badge(props: dict) -> str:
    c = props.get("confidence")
    if c is None:
        return ""
    pct = int(c * 100)
    if pct >= 90:
        tier = "high"
    elif pct >= 60:
        tier = "medium"
    else:
        tier = "low"
    return f"  [confidence: {pct}% / {tier}]"


def _temporal_badge(props: dict) -> str:
    now = _now_iso()[:10]
    parts = []
    vf = props.get("valid_from")
    vt = props.get("valid_to")
    if vf:
        vf_short = vf[:10]
        parts.append(f"from {vf_short}")
        # Add staleness hint for old knowledge
        try:
            from datetime import datetime, timezone
            age_days = (datetime.now(timezone.utc) - datetime.fromisoformat(vf)).days
            if age_days > 365:
                parts.append(f"⚠ {age_days // 365}y old")
            elif age_days > 90:
                parts.append(f"⚠ {age_days}d old")
        except Exception:
            pass
    if vt:
        vt_short = vt[:10]
        parts.append(f"expires {vt_short}")
        # Warn if expiring soon
        if vt_short <= now:
            parts.append("EXPIRED")
        elif vt[:10] <= (now[:8] + str(int(now[8:10]) + 7).zfill(2))[:10]:
            parts.append("expiring soon")
    return f"  [{', '.join(parts)}]" if parts else ""


def resolve_belief_conflicts(records: list) -> list:
    """
    For records with belief contradictions, pick the winning belief by:
    1. Highest confidence
    2. Most recent valid_from as tiebreaker
    Sets _resolved_belief and _resolution_reason on each conflicting record.
    """
    for rec in records:
        contradictions = [c for c in (rec.get("contradictions") or []) if c]
        if not contradictions:
            continue
        node = rec.get("node", {})
        candidates = [node] + contradictions
        def _sort_key(b):
            conf = float(b.get("confidence") or 0)
            vf   = str(b.get("valid_from") or "")
            return (conf, vf)
        winner = max(candidates, key=_sort_key)
        conf_pct = int(float(winner.get("confidence") or 0) * 100)
        vf = str(winner.get("valid_from") or "")[:10]
        rec["_resolved_belief"] = winner
        rec["_resolution_reason"] = (
            f"Selected belief with confidence={conf_pct}%"
            + (f", valid_from={vf}" if vf else "")
            + f" (resolved from {len(candidates)} candidates)"
        )
    return records


def build_gcor_context(intent: str, qdrant_hits: list, graph_records: list, provenance: str = "") -> str:
    records = _reflection_fallback(qdrant_hits, graph_records)
    if not records:
        return ""

    confidence_note = (
        f"confidence ≥ {int(CONFIDENCE_THRESHOLD*100)}% "
        if CONFIDENCE_THRESHOLD > 0 else ""
    )
    agent_note = f"agent partition: {AGENT_ID} " if AGENT_ID else ""

    now = _now_iso()
    lines = [
        f"## Retrieved Context  [intent: {intent} | "
        f"{len(qdrant_hits)} semantic candidates → {len(records)} graph records"
        + (f" | {confidence_note}{agent_note}".rstrip() if (confidence_note or agent_note) else "")
        + "]",
        "",
        f"Current time: {now[:19]}Z. "
        "Neo4j is the primary source of truth. Qdrant identified the entry points. "
        "All results have been temporally filtered — only currently valid knowledge "
        "is shown. Prefer more recent sources when information conflicts.",
    ]
    if provenance:
        lines += ["", f"**Retrieval Provenance:** {provenance}"]
    lines.append("")

    for rec in records:
        node   = rec.get("node") or {}
        labels = rec.get("labels") or []
        doc    = rec.get("document") or {}

        node_type = labels[0] if labels else "Node"
        content   = (
            node.get("text") or node.get("content") or
            node.get("description") or node.get("title") or
            node.get("subject") or str(node)
        )[:600]

        conf_badge    = _confidence_badge(node)
        temporal_badge = _temporal_badge(node)
        lines.append(f"### [{node_type}]{conf_badge}{temporal_badge}")
        lines.append(content)

        if doc.get("title") or doc.get("source"):
            lines.append(f"Source: {doc.get('title') or doc.get('source', '')}")

        # Reasoning trace for Inference nodes
        if node.get("reasoning_trace"):
            lines.append(f"Reasoning: {node['reasoning_trace'][:300]}")

        # Intent-specific extras
        if intent == "planning":
            deps = [d for d in (rec.get("dependencies") or []) if d]
            if deps:
                lines.append("Dependencies: " + " → ".join(
                    d.get("description") or d.get("title") or str(d) for d in deps
                ))
            blocking = [g for g in (rec.get("blocking_goals") or []) if g]
            if blocking:
                lines.append("Blocked by: " + ", ".join(
                    g.get("description") or str(g) for g in blocking
                ))
            inferences = [i for i in (rec.get("supporting_inferences") or []) if i]
            if inferences:
                lines.append("Supporting inferences: " + "; ".join(
                    f"{i.get('text', '')[:80]} ({int(i.get('confidence', 0)*100)}%)"
                    for i in inferences
                ))

        elif intent == "dependency":
            chain = [x for x in (rec.get("dependency_chain") or []) if x]
            if len(chain) > 1:
                lines.append("Dependency path: " + " → ".join(
                    x.get("description") or x.get("title") or str(x) for x in chain
                ))

        elif intent == "memory":
            mems = [m for m in (rec.get("memories") or []) if m]
            if mems:
                lines.append("Memories: " + "; ".join(
                    f"{m.get('content', '')[:80]}{_confidence_badge(m)}"
                    for m in mems[:3]
                ))
            beliefs = [b for b in (rec.get("beliefs") or []) if b]
            if beliefs:
                lines.append("Beliefs: " + "; ".join(
                    f"{b.get('content', '')[:80]}{_confidence_badge(b)}"
                    for b in beliefs[:3]
                ))
            evts = [e for e in (rec.get("events") or []) if e]
            if evts:
                lines.append("Events: " + "; ".join(
                    f"{e.get('timestamp', '')[:10]} {e.get('description', '')}"
                    for e in evts[:3]
                ))

        elif intent == "inference":
            srcs = [s for s in (rec.get("sources") or []) if s]
            if srcs:
                lines.append("Derived from: " + "; ".join(
                    (s.get("text") or s.get("content") or str(s))[:80]
                    for s in srcs[:3]
                ))
            supports = [t for t in (rec.get("supports") or []) if t]
            if supports:
                lines.append("Supports: " + "; ".join(
                    (t.get("description") or t.get("content") or str(t))[:80]
                    for t in supports[:3]
                ))
            downstream = [d for d in (rec.get("downstream_inferences") or []) if d]
            if downstream:
                lines.append("Downstream inferences: " + "; ".join(
                    f"{d.get('text', '')[:60]}{_confidence_badge(d)}"
                    for d in downstream[:2]
                ))

        elif intent == "belief":
            contradictions = [c for c in (rec.get("contradictions") or []) if c]
            if contradictions:
                lines.append("Contradicts: " + "; ".join(
                    f"{c.get('content', '')[:80]}{_confidence_badge(c)}"
                    for c in contradictions[:2]
                ))
            inferences = [i for i in (rec.get("supporting_inferences") or []) if i]
            if inferences:
                lines.append("Supported by: " + "; ".join(
                    f"{i.get('text', '')[:60]}{_confidence_badge(i)}"
                    for i in inferences[:2]
                ))
            agent = rec.get("agent") or {}
            if agent.get("id"):
                lines.append(f"Held by agent: {agent['id']}")
            if rec.get("_resolution_reason"):
                lines.append(f"⚖️ Conflict resolution: {rec['_resolution_reason']}")

        elif intent == "semantic":
            concepts = [c for c in (rec.get("concepts") or []) if c]
            if concepts:
                lines.append("Concepts: " + ", ".join(c.get("name", "") for c in concepts))
            neighbors = [n for n in (rec.get("neighbors") or []) if n]
            if neighbors:
                lines.append("Adjacent chunks: " + " | ".join(
                    n.get("text", "")[:100] for n in neighbors
                ))

        else:  # factual
            related = [r for r in (rec.get("related") or []) if r and r.get("props")]
            if related:
                lines.append("Related: " + "; ".join(
                    f"{r['rel']} → {r['props'].get('name') or r['props'].get('text', '')[:80]}"
                    for r in related[:4]
                ))

        lines.append("")

    lines.append(
        "Instruction: use the retrieved context above when relevant. "
        "Respect the confidence scores — lower-confidence information should be "
        "presented with appropriate uncertainty. If the context is insufficient, "
        "apply general knowledge."
    )
    return "\n".join(lines)


def inject_context(messages: list, context: str) -> list:
    if not context:
        return messages
    result, merged = [], False
    for m in messages:
        if m.get("role") == "system" and not merged:
            result.append({"role": "system",
                           "content": context + "\n\n---\n\n" + m.get("content", "")})
            merged = True
        else:
            result.append(m)
    if not merged:
        result = [{"role": "system", "content": context}] + messages
    return result


def last_user_text(messages: list) -> str:
    for m in reversed(messages):
        if m.get("role") == "user":
            content = m.get("content", "")
            if isinstance(content, str):
                return content
            if isinstance(content, list):
                return " ".join(
                    p.get("text", "") for p in content
                    if isinstance(p, dict) and p.get("type") == "text"
                )
    return ""


# ── LLM routing ───────────────────────────────────────────────────────────────

def _resolve_model(model: str) -> tuple:
    if model == "openclaw":
        if LLM_BACKEND == "openclaw":
            return "openclaw", "openclaw"
        if LLM_BACKEND == "copilot":
            return COPILOT_CHAT_MODEL, "copilot"
        if LLM_BACKEND == "ollama":
            return OLLAMA_MODEL, "ollama"
        if LLM_BACKEND == "openai":
            return OPENAI_CHAT_MODEL, "openai"
        return ANTHROPIC_CHAT_MODEL, "anthropic"
    if model.startswith("claude-"):
        return model, "anthropic"
    if model.startswith(("gpt-", "o1", "o3")):
        return model, "openai"
    if LLM_BACKEND == "openclaw":
        return "openclaw", "openclaw"
    if LLM_BACKEND == "copilot":
        return COPILOT_CHAT_MODEL, "copilot"
    if LLM_BACKEND == "anthropic":
        return ANTHROPIC_CHAT_MODEL, "anthropic"
    if LLM_BACKEND == "ollama":
        return OLLAMA_MODEL, "ollama"
    return OPENAI_CHAT_MODEL, "openai"


async def call_ollama(body: dict):
    """Forward chat completions to the local Ollama OpenAI-compatible endpoint."""
    url = f"{OLLAMA_BASE_URL}/v1/chat/completions"
    body = {**body, "model": body.get("model", OLLAMA_MODEL)}
    if body.get("stream"):
        async def stream():
            try:
                METRIC_LLM_REQUESTS.labels(backend="ollama", status="started").inc()
                async with httpx.AsyncClient(timeout=300) as c:
                    async with c.stream("POST", url, json=body) as resp:
                        if resp.status_code != 200:
                            err = await resp.aread()
                            logger.error("Ollama %s: %s", resp.status_code, err[:200])
                            METRIC_LLM_REQUESTS.labels(backend="ollama", status="error").inc()
                            yield f'data: {{"error": "Ollama {resp.status_code}"}}\n\n'.encode()
                            return
                        METRIC_LLM_REQUESTS.labels(backend="ollama", status="success").inc()
                        async for chunk in resp.aiter_bytes():
                            yield chunk
            except Exception as exc:
                logger.error("Ollama stream error: %s", exc)
                METRIC_LLM_REQUESTS.labels(backend="ollama", status="exception").inc()
                yield b'data: {"error": "proxy error"}\n\n'
        return StreamingResponse(stream(), media_type="text/event-stream")
    else:
        t0 = time.monotonic()
        try:
            async with httpx.AsyncClient(timeout=300) as c:
                resp = await c.post(url, json=body)
                resp.raise_for_status()
            METRIC_LLM_REQUESTS.labels(backend="ollama", status="success").inc()
            METRIC_LLM_DURATION.labels(backend="ollama").observe(time.monotonic() - t0)
            return JSONResponse(content=resp.json(), status_code=resp.status_code)
        except Exception:
            METRIC_LLM_REQUESTS.labels(backend="ollama", status="error").inc()
            raise


async def call_openclaw(body: dict):
    """Forward chat completions to the openclaw gateway OpenAI-compatible endpoint."""
    url = f"{OPENCLAW_BASE_URL}/chat/completions"
    headers = {"Authorization": f"Bearer {OPENCLAW_GATEWAY_TOKEN}", "Content-Type": "application/json"}
    body = {**body, "model": body.get("model", "default")}
    if body.get("stream"):
        async def stream():
            try:
                METRIC_LLM_REQUESTS.labels(backend="openclaw", status="started").inc()
                async with httpx.AsyncClient(timeout=300) as c:
                    async with c.stream("POST", url, json=body, headers=headers) as resp:
                        if resp.status_code != 200:
                            err = await resp.aread()
                            logger.error("OpenClaw %s: %s", resp.status_code, err[:200])
                            METRIC_LLM_REQUESTS.labels(backend="openclaw", status="error").inc()
                            yield f'data: {{"error": "OpenClaw {resp.status_code}"}}\n\n'.encode()
                            return
                        METRIC_LLM_REQUESTS.labels(backend="openclaw", status="success").inc()
                        async for chunk in resp.aiter_bytes():
                            yield chunk
            except Exception as exc:
                logger.error("OpenClaw stream error: %s", exc)
                METRIC_LLM_REQUESTS.labels(backend="openclaw", status="exception").inc()
                yield b'data: {"error": "proxy error"}\n\n'
        return StreamingResponse(stream(), media_type="text/event-stream")
    else:
        t0 = time.monotonic()
        try:
            async with httpx.AsyncClient(timeout=300) as c:
                resp = await c.post(url, json=body, headers=headers)
                resp.raise_for_status()
            METRIC_LLM_REQUESTS.labels(backend="openclaw", status="success").inc()
            METRIC_LLM_DURATION.labels(backend="openclaw").observe(time.monotonic() - t0)
            return JSONResponse(content=resp.json(), status_code=resp.status_code)
        except Exception:
            METRIC_LLM_REQUESTS.labels(backend="openclaw", status="error").inc()
            raise


async def call_copilot(body: dict):
    """Forward chat completions to GitHub Copilot (OpenAI-compatible, gpt-4.1)."""
    url = f"{COPILOT_BASE_URL}/chat/completions"
    headers = {
        "Authorization": f"Bearer {COPILOT_API_KEY}",
        "Content-Type": "application/json",
        "Copilot-Integration-Id": "vscode-chat",
        "Editor-Version": "vscode/1.90.0",
        "Editor-Plugin-Version": "copilot-chat/0.22.4",
        "Openai-Intent": "conversation-panel",
        "x-github-api-version": "2023-07-07",
    }
    body = {**body, "model": body.get("model", COPILOT_CHAT_MODEL)}
    if body.get("stream"):
        async def stream():
            for attempt in range(LLM_MAX_RETRIES + 1):
                try:
                    METRIC_LLM_REQUESTS.labels(backend="copilot", status="started").inc()
                    async with httpx.AsyncClient(timeout=120) as c:
                        async with c.stream("POST", url, json=body, headers=headers) as resp:
                            if resp.status_code != 200:
                                err = await resp.aread()
                                if _is_retryable(resp.status_code) and attempt < LLM_MAX_RETRIES:
                                    ra = resp.headers.get("retry-after")
                                    wait = float(ra) if ra else (LLM_RETRY_BASE * (2 ** attempt) + random.uniform(0, 0.5))
                                    logger.warning("Copilot %s (attempt %d/%d), retrying in %.1fs",
                                                   resp.status_code, attempt + 1, LLM_MAX_RETRIES + 1, wait)
                                    await asyncio.sleep(wait)
                                    continue
                                logger.error("Copilot %s: %s", resp.status_code, err[:200])
                                METRIC_LLM_REQUESTS.labels(backend="copilot", status="error").inc()
                                yield f'data: {{"error": "Copilot {resp.status_code}"}}\n\n'.encode()
                                return
                            METRIC_LLM_REQUESTS.labels(backend="copilot", status="success").inc()
                            async for chunk in resp.aiter_bytes():
                                yield chunk
                            return
                except Exception as exc:
                    if attempt < LLM_MAX_RETRIES:
                        wait = LLM_RETRY_BASE * (2 ** attempt) + random.uniform(0, 0.5)
                        logger.warning("Copilot stream error (attempt %d/%d): %s, retrying in %.1fs",
                                       attempt + 1, LLM_MAX_RETRIES + 1, exc, wait)
                        await asyncio.sleep(wait)
                        continue
                    logger.error("Copilot stream error: %s", exc)
                    METRIC_LLM_REQUESTS.labels(backend="copilot", status="exception").inc()
                    yield b'data: {"error": "proxy error"}\n\n'
                    return
        return StreamingResponse(stream(), media_type="text/event-stream")
    else:
        t0 = time.monotonic()
        for attempt in range(LLM_MAX_RETRIES + 1):
            try:
                async with httpx.AsyncClient(timeout=120) as c:
                    resp = await c.post(url, json=body, headers=headers)
                    if _is_retryable(resp.status_code) and attempt < LLM_MAX_RETRIES:
                        ra = resp.headers.get("retry-after")
                        wait = float(ra) if ra else (LLM_RETRY_BASE * (2 ** attempt) + random.uniform(0, 0.5))
                        logger.warning("Copilot %s (attempt %d/%d), retrying in %.1fs",
                                       resp.status_code, attempt + 1, LLM_MAX_RETRIES + 1, wait)
                        await asyncio.sleep(wait)
                        continue
                    resp.raise_for_status()
                METRIC_LLM_REQUESTS.labels(backend="copilot", status="success").inc()
                METRIC_LLM_DURATION.labels(backend="copilot").observe(time.monotonic() - t0)
                return JSONResponse(content=resp.json(), status_code=resp.status_code)
            except httpx.HTTPStatusError:
                METRIC_LLM_REQUESTS.labels(backend="copilot", status="error").inc()
                raise
            except Exception:
                if attempt < LLM_MAX_RETRIES:
                    wait = LLM_RETRY_BASE * (2 ** attempt) + random.uniform(0, 0.5)
                    logger.warning("Copilot non-stream error (attempt %d/%d), retrying in %.1fs",
                                   attempt + 1, LLM_MAX_RETRIES + 1, wait)
                    await asyncio.sleep(wait)
                    continue
                METRIC_LLM_REQUESTS.labels(backend="copilot", status="error").inc()
                raise


async def call_openai(body: dict):
    headers = {"Authorization": f"Bearer {OPENAI_API_KEY}", "Content-Type": "application/json"}
    if body.get("stream"):
        async def stream():
            for attempt in range(LLM_MAX_RETRIES + 1):
                try:
                    METRIC_LLM_REQUESTS.labels(backend="openai", status="started").inc()
                    async with httpx.AsyncClient(timeout=120) as c:
                        async with c.stream("POST", "https://api.openai.com/v1/chat/completions",
                                            json=body, headers=headers) as resp:
                            if resp.status_code != 200:
                                err = await resp.aread()
                                if _is_retryable(resp.status_code) and attempt < LLM_MAX_RETRIES:
                                    ra = resp.headers.get("retry-after")
                                    wait = float(ra) if ra else (LLM_RETRY_BASE * (2 ** attempt) + random.uniform(0, 0.5))
                                    logger.warning("OpenAI %s (attempt %d/%d), retrying in %.1fs",
                                                   resp.status_code, attempt + 1, LLM_MAX_RETRIES + 1, wait)
                                    await asyncio.sleep(wait)
                                    continue
                                logger.error("OpenAI %s: %s", resp.status_code, err[:200])
                                METRIC_LLM_REQUESTS.labels(backend="openai", status="error").inc()
                                yield f'data: {{"error": "OpenAI {resp.status_code}"}}\n\n'.encode()
                                return
                            METRIC_LLM_REQUESTS.labels(backend="openai", status="success").inc()
                            async for chunk in resp.aiter_bytes():
                                yield chunk
                            return
                except Exception as exc:
                    if attempt < LLM_MAX_RETRIES:
                        wait = LLM_RETRY_BASE * (2 ** attempt) + random.uniform(0, 0.5)
                        logger.warning("OpenAI stream error (attempt %d/%d): %s, retrying in %.1fs",
                                       attempt + 1, LLM_MAX_RETRIES + 1, exc, wait)
                        await asyncio.sleep(wait)
                        continue
                    logger.error("OpenAI stream error: %s", exc)
                    METRIC_LLM_REQUESTS.labels(backend="openai", status="exception").inc()
                    yield b'data: {"error": "proxy error"}\n\n'
                    return
        return StreamingResponse(stream(), media_type="text/event-stream")
    else:
        t0 = time.monotonic()
        for attempt in range(LLM_MAX_RETRIES + 1):
            try:
                async with httpx.AsyncClient(timeout=120) as c:
                    resp = await c.post("https://api.openai.com/v1/chat/completions",
                                        json=body, headers=headers)
                    if _is_retryable(resp.status_code) and attempt < LLM_MAX_RETRIES:
                        ra = resp.headers.get("retry-after")
                        wait = float(ra) if ra else (LLM_RETRY_BASE * (2 ** attempt) + random.uniform(0, 0.5))
                        logger.warning("OpenAI %s (attempt %d/%d), retrying in %.1fs",
                                       resp.status_code, attempt + 1, LLM_MAX_RETRIES + 1, wait)
                        await asyncio.sleep(wait)
                        continue
                    resp.raise_for_status()
                METRIC_LLM_REQUESTS.labels(backend="openai", status="success").inc()
                METRIC_LLM_DURATION.labels(backend="openai").observe(time.monotonic() - t0)
                return JSONResponse(content=resp.json(), status_code=resp.status_code)
            except httpx.HTTPStatusError:
                METRIC_LLM_REQUESTS.labels(backend="openai", status="error").inc()
                raise
            except Exception:
                if attempt < LLM_MAX_RETRIES:
                    wait = LLM_RETRY_BASE * (2 ** attempt) + random.uniform(0, 0.5)
                    logger.warning("OpenAI non-stream error (attempt %d/%d), retrying in %.1fs",
                                   attempt + 1, LLM_MAX_RETRIES + 1, wait)
                    await asyncio.sleep(wait)
                    continue
                METRIC_LLM_REQUESTS.labels(backend="openai", status="error").inc()
                raise


def _is_retryable(status_code: int) -> bool:
    """Return True for status codes that warrant a retry (rate-limit, overloaded, server error)."""
    return status_code in (429, 529) or status_code >= 500


async def call_anthropic(body: dict):
    messages   = body.get("messages", [])
    system_txt = next((m["content"] for m in messages if m.get("role") == "system"), None)
    user_msgs  = [m for m in messages if m.get("role") != "system"]
    ant_body   = {"model": body.get("model", ANTHROPIC_CHAT_MODEL),
                  "max_tokens": body.get("max_tokens", 4096), "messages": user_msgs}
    if system_txt:
        ant_body["system"] = system_txt
    headers = {"x-api-key": ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01",
               "Content-Type": "application/json"}

    if body.get("stream"):
        ant_body["stream"] = True
        async def stream():
            last_err = None
            for attempt in range(LLM_MAX_RETRIES + 1):
                try:
                    METRIC_LLM_REQUESTS.labels(backend="anthropic", status="started").inc()
                    async with httpx.AsyncClient(timeout=120) as c:
                        async with c.stream("POST", "https://api.anthropic.com/v1/messages",
                                            json=ant_body, headers=headers) as resp:
                            if resp.status_code != 200:
                                err = await resp.aread()
                                last_err = f"Anthropic {resp.status_code}"
                                if _is_retryable(resp.status_code) and attempt < LLM_MAX_RETRIES:
                                    ra = resp.headers.get("retry-after")
                                    wait = float(ra) if ra else (LLM_RETRY_BASE * (2 ** attempt) + random.uniform(0, 0.5))
                                    logger.warning("Anthropic %s (attempt %d/%d), retrying in %.1fs",
                                                   resp.status_code, attempt + 1, LLM_MAX_RETRIES + 1, wait)
                                    await asyncio.sleep(wait)
                                    continue
                                logger.error("Anthropic %s: %s", resp.status_code, err[:200])
                                METRIC_LLM_REQUESTS.labels(backend="anthropic", status="error").inc()
                                yield f'data: {{"error": "Anthropic {resp.status_code}"}}\n\n'.encode()
                                return
                            METRIC_LLM_REQUESTS.labels(backend="anthropic", status="success").inc()
                            async for line in resp.aiter_lines():
                                if not line.startswith("data: "):
                                    continue
                                raw = line[6:]
                                if raw == "[DONE]":
                                    yield b"data: [DONE]\n\n"; continue
                                try:
                                    evt = json.loads(raw)
                                    if evt.get("type") == "content_block_delta":
                                        txt = evt.get("delta", {}).get("text", "")
                                        yield f'data: {json.dumps({"choices": [{"delta": {"content": txt}, "finish_reason": None}]})}\n\n'.encode()
                                    elif evt.get("type") == "message_stop":
                                        yield b"data: [DONE]\n\n"
                                except json.JSONDecodeError:
                                    pass
                            return
                except Exception as exc:
                    last_err = str(exc)
                    if attempt < LLM_MAX_RETRIES:
                        wait = LLM_RETRY_BASE * (2 ** attempt) + random.uniform(0, 0.5)
                        logger.warning("Anthropic stream error (attempt %d/%d): %s, retrying in %.1fs",
                                       attempt + 1, LLM_MAX_RETRIES + 1, exc, wait)
                        await asyncio.sleep(wait)
                        continue
                    logger.error("Anthropic stream error: %s", exc)
                    METRIC_LLM_REQUESTS.labels(backend="anthropic", status="exception").inc()
                    yield b'data: {"error": "proxy error"}\n\n'
                    return
        return StreamingResponse(stream(), media_type="text/event-stream")
    else:
        t0 = time.monotonic()
        for attempt in range(LLM_MAX_RETRIES + 1):
            try:
                async with httpx.AsyncClient(timeout=120) as c:
                    resp = await c.post("https://api.anthropic.com/v1/messages",
                                        json=ant_body, headers=headers)
                    if _is_retryable(resp.status_code) and attempt < LLM_MAX_RETRIES:
                        ra = resp.headers.get("retry-after")
                        wait = float(ra) if ra else (LLM_RETRY_BASE * (2 ** attempt) + random.uniform(0, 0.5))
                        logger.warning("Anthropic %s (attempt %d/%d), retrying in %.1fs",
                                       resp.status_code, attempt + 1, LLM_MAX_RETRIES + 1, wait)
                        await asyncio.sleep(wait)
                        continue
                    resp.raise_for_status()
                METRIC_LLM_REQUESTS.labels(backend="anthropic", status="success").inc()
                METRIC_LLM_DURATION.labels(backend="anthropic").observe(time.monotonic() - t0)
                ant = resp.json()
                usage = ant.get("usage", {})
                return JSONResponse(content={
                    "id": ant.get("id", ""), "object": "chat.completion",
                    "model": ant.get("model", ant_body["model"]),
                    "choices": [{"index": 0,
                                 "message": {"role": "assistant",
                                             "content": ant.get("content", [{}])[0].get("text", "")},
                                 "finish_reason": ant.get("stop_reason", "stop")}],
                    "usage": {"prompt_tokens":    usage.get("input_tokens", 0),
                              "completion_tokens": usage.get("output_tokens", 0),
                              "total_tokens":      usage.get("input_tokens", 0) + usage.get("output_tokens", 0)},
                })
            except httpx.HTTPStatusError:
                METRIC_LLM_REQUESTS.labels(backend="anthropic", status="error").inc()
                raise
            except Exception:
                if attempt < LLM_MAX_RETRIES:
                    wait = LLM_RETRY_BASE * (2 ** attempt) + random.uniform(0, 0.5)
                    logger.warning("Anthropic non-stream error (attempt %d/%d), retrying in %.1fs",
                                   attempt + 1, LLM_MAX_RETRIES + 1, wait)
                    await asyncio.sleep(wait)
                    continue
                METRIC_LLM_REQUESTS.labels(backend="anthropic", status="error").inc()
                raise


# ── Routes ─────────────────────────────────────────────────────────────────────

@app.get("/v1/models")
async def list_models():
    models = []
    # Pull Ollama models
    if LLM_BACKEND == "ollama" or EMBEDDING_BACKEND == "ollama":
        try:
            async with httpx.AsyncClient(timeout=5) as c:
                resp = await c.get(f"{OLLAMA_BASE_URL}/api/tags")
                if resp.status_code == 200:
                    for m in resp.json().get("models", []):
                        models.append({"id": m["name"], "object": "model", "owned_by": "ollama"})
        except Exception as exc:
            logger.warning("Could not fetch Ollama models: %s", exc)
    # Pull OpenAI models
    if OPENAI_API_KEY:
        try:
            async with httpx.AsyncClient(timeout=10) as c:
                resp = await c.get("https://api.openai.com/v1/models",
                                   headers={"Authorization": f"Bearer {OPENAI_API_KEY}"})
                if resp.status_code == 200:
                    models += resp.json().get("data", [])
        except Exception as exc:
            logger.warning("Could not fetch OpenAI models: %s", exc)
    if not models:
        models = [
            {"id": OPENAI_CHAT_MODEL,    "object": "model", "owned_by": "openai"},
            {"id": ANTHROPIC_CHAT_MODEL, "object": "model", "owned_by": "anthropic"},
            {"id": OLLAMA_MODEL,         "object": "model", "owned_by": "ollama"},
        ]
    models = [{"id": "openclaw", "object": "model", "created": 1700000000, "owned_by": "eedgeai"}] + models
    return {"object": "list", "data": models}


@app.post("/v1/embeddings")
async def embeddings(request: Request):
    """Proxy embedding requests to the configured backend (for OpenWebUI's built-in RAG)."""
    body = await request.json()
    try:
        if EMBEDDING_BACKEND == "ollama":
            async with httpx.AsyncClient(timeout=60) as c:
                resp = await c.post(f"{OLLAMA_BASE_URL}/v1/embeddings", json=body)
                resp.raise_for_status()
                return JSONResponse(content=resp.json())
        if not OPENAI_API_KEY:
            raise HTTPException(status_code=503, detail="OPENAI_API_KEY not configured")
        async with httpx.AsyncClient(timeout=30) as c:
            resp = await c.post(
                "https://api.openai.com/v1/embeddings",
                headers={"Authorization": f"Bearer {OPENAI_API_KEY}",
                         "Content-Type": "application/json"},
                json=body,
            )
            resp.raise_for_status()
            return JSONResponse(content=resp.json())
    except httpx.HTTPStatusError as exc:
        raise HTTPException(status_code=exc.response.status_code,
                            detail="Embeddings upstream error")
    except Exception as exc:
        logger.error("Embeddings failed: %s", exc)
        raise HTTPException(status_code=502, detail="Embeddings request failed")


@app.post("/v1/chat/completions")
async def chat_completions(request: Request):
    body     = await request.json()
    messages = list(body.get("messages", []))

    rag_intent   = "none"
    rag_fallback = "disabled"
    rag_t0       = time.monotonic()

    if ENABLE_RAG and messages:
        query = last_user_text(messages)

        if query:
            # ── Step 1: Intent ──────────────────────────────────────────────
            intent     = await classify_intent(query)
            rag_intent = intent
            logger.info("GCOR intent: %s", intent)

            # ── Step 2a: Semantic (Qdrant → element IDs) ────────────────────
            vector      = await embed_text(query)
            raw_hits    = await qdrant_search(vector) if vector else []
            qdrant_hits = _filter_hits(raw_hits)
            element_ids = [
                h["payload"]["neo4j_element_id"]
                for h in qdrant_hits
                if h.get("payload", {}).get("neo4j_element_id")
            ]
            METRIC_QDRANT_HITS.observe(len(qdrant_hits))
            logger.info(
                "Qdrant: %d raw → %d after cognitive filters, %d with neo4j_element_id",
                len(raw_hits), len(qdrant_hits), len(element_ids),
            )

            # ── Step 2b: Structural (Neo4j graph expansion) ─────────────────
            graph_records = await neo4j_expand(element_ids, intent, query)
            METRIC_NEO4J_RECORDS.observe(len(graph_records))
            logger.info("Neo4j: %d graph records expanded", len(graph_records))

            # ── Step 2c: Belief conflict resolution ─────────────────────────
            graph_records = resolve_belief_conflicts(graph_records)

            # ── Step 2d: Provenance chain ────────────────────────────────────
            provenance = await build_provenance_chain(element_ids)
            if provenance:
                logger.info("Provenance chain: %s", provenance[:120])

            # ── Step 3: Reflection + context build ──────────────────────────
            context  = build_gcor_context(intent, qdrant_hits, graph_records, provenance=provenance)

            # Determine fallback mode for metrics
            if graph_records:
                rag_fallback = "full_gcor"
            elif qdrant_hits:
                rag_fallback = "qdrant_text"
            else:
                rag_fallback = "llm_only"

            messages = inject_context(messages, context)
            body     = {**body, "messages": messages}

        METRIC_RAG_REQUESTS.labels(intent=rag_intent, fallback_mode=rag_fallback).inc()
        METRIC_RAG_DURATION.observe(time.monotonic() - rag_t0)

    # ── Route to LLM ──────────────────────────────────────────────────────────
    model, backend = _resolve_model(body.get("model", "openclaw"))
    body = {**body, "model": model}
    try:
        if backend == "openclaw":
            if body.get("stream") and COPILOT_API_KEY:
                # Streaming: build inline fallback generator (exceptions inside
                # StreamingResponse generators are not catchable from outside).
                oc_url = f"{OPENCLAW_BASE_URL}/chat/completions"
                oc_hdrs = {"Authorization": f"Bearer {OPENCLAW_GATEWAY_TOKEN}",
                           "Content-Type": "application/json"}
                cp_url = f"{COPILOT_BASE_URL}/chat/completions"
                cp_hdrs = {
                    "Authorization": f"Bearer {COPILOT_API_KEY}",
                    "Content-Type": "application/json",
                    "Copilot-Integration-Id": "vscode-chat",
                    "Editor-Version": "vscode/1.90.0",
                    "Editor-Plugin-Version": "copilot-chat/0.22.4",
                    "Openai-Intent": "conversation-panel",
                    "x-github-api-version": "2023-07-07",
                }
                cp_body = {**body, "model": COPILOT_CHAT_MODEL}

                async def _stream_openclaw_or_copilot():
                    oc_ok = False
                    try:
                        METRIC_LLM_REQUESTS.labels(backend="openclaw", status="started").inc()
                        async with httpx.AsyncClient(timeout=300) as c:
                            async with c.stream("POST", oc_url, json=body, headers=oc_hdrs) as resp:
                                if resp.status_code == 200:
                                    METRIC_LLM_REQUESTS.labels(backend="openclaw", status="success").inc()
                                    oc_ok = True
                                    async for chunk in resp.aiter_bytes():
                                        yield chunk
                                    return
                                err = await resp.aread()
                                logger.warning("OpenClaw %s, falling back to Copilot gpt-4.1: %s",
                                               resp.status_code, err[:200])
                                METRIC_LLM_REQUESTS.labels(backend="openclaw", status="error").inc()
                    except Exception as exc:
                        logger.warning("OpenClaw stream error, falling back to Copilot: %s", exc)
                        METRIC_LLM_REQUESTS.labels(backend="openclaw", status="exception").inc()

                    if oc_ok:
                        return

                    logger.info("Streaming via GitHub Copilot gpt-4.1 fallback")
                    METRIC_LLM_REQUESTS.labels(backend="copilot", status="started").inc()
                    try:
                        async with httpx.AsyncClient(timeout=120) as c:
                            async with c.stream("POST", cp_url, json=cp_body, headers=cp_hdrs) as resp:
                                if resp.status_code != 200:
                                    err = await resp.aread()
                                    logger.error("Copilot fallback %s: %s", resp.status_code, err[:200])
                                    METRIC_LLM_REQUESTS.labels(backend="copilot", status="error").inc()
                                    yield f'data: {{"error": "All backends failed ({resp.status_code})"}}\n\n'.encode()
                                    return
                                METRIC_LLM_REQUESTS.labels(backend="copilot", status="success").inc()
                                async for chunk in resp.aiter_bytes():
                                    yield chunk
                    except Exception as exc:
                        logger.error("Copilot fallback stream error: %s", exc)
                        METRIC_LLM_REQUESTS.labels(backend="copilot", status="exception").inc()
                        yield b'data: {"error": "proxy error"}\n\n'

                return StreamingResponse(_stream_openclaw_or_copilot(), media_type="text/event-stream")
            else:
                # Non-streaming (or no Copilot key) — exceptions propagate to outer handler
                try:
                    return await call_openclaw(body)
                except Exception as exc:
                    if COPILOT_API_KEY:
                        logger.warning("OpenClaw failed (%s), falling back to Copilot gpt-4.1", exc)
                        METRIC_LLM_REQUESTS.labels(backend="copilot", status="fallback").inc()
                        return await call_copilot({**body, "model": COPILOT_CHAT_MODEL})
                    raise
        elif backend == "copilot" and COPILOT_API_KEY:
            return await call_copilot(body)
        elif backend == "ollama":
            return await call_ollama(body)
        elif backend == "anthropic" and ANTHROPIC_API_KEY:
            return await call_anthropic(body)
        elif OPENAI_API_KEY:
            return await call_openai(body)
        else:
            raise HTTPException(status_code=503, detail="No LLM backend configured")
    except httpx.HTTPStatusError as exc:
        logger.error("LLM upstream error: %s", exc)
        raise HTTPException(status_code=exc.response.status_code, detail="LLM upstream error")
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("Chat completion failed: %s", exc)
        raise HTTPException(status_code=502, detail="Chat completion failed")


_TEMPLATES = os.path.join(os.path.dirname(__file__), "templates")


# ── Knowledge UI ───────────────────────────────────────────────────────────────

@app.get("/", response_class=RedirectResponse, status_code=302)
async def root():
    return "/knowledge"


@app.get("/openclaw", response_class=RedirectResponse, status_code=302)
async def openclaw_redirect():
    token = OPENCLAW_GATEWAY_TOKEN
    base  = "http://127.0.0.1:18799"
    if token:
        return f"{base}/#token={token}"
    return base


@app.get("/knowledge", response_class=HTMLResponse)
async def knowledge_ui():
    with open(os.path.join(_TEMPLATES, "knowledge.html"), encoding="utf-8") as f:
        return f.read()


@app.get("/api/collections")
async def api_collections():
    """List all Qdrant collections with Neo4j document stats."""
    # Fetch Qdrant collections
    collections = []
    try:
        async with httpx.AsyncClient(timeout=10) as c:
            resp = await c.get(f"{QDRANT_URL}/collections")
            resp.raise_for_status()
            for col in resp.json().get("result", {}).get("collections", []):
                name = col["name"]
                info_resp = await c.get(f"{QDRANT_URL}/collections/{name}")
                info = info_resp.json().get("result", {})
                collections.append({
                    "name":            name,
                    "points_count":    info.get("points_count", 0),
                    "doc_count":       0,
                    "recent_docs":     [],
                    "embedding_model": EMBEDDING_MODEL,
                })
    except Exception as exc:
        logger.warning("Qdrant collection fetch failed: %s", exc)

    # If the configured collection is missing from Qdrant, still show it
    names = {c["name"] for c in collections}
    if QDRANT_COLLECTION not in names:
        collections.insert(0, {
            "name":            QDRANT_COLLECTION,
            "points_count":    0,
            "doc_count":       0,
            "recent_docs":     [],
            "embedding_model": EMBEDDING_MODEL,
        })

    # Enrich with Neo4j document counts
    driver = AsyncGraphDatabase.driver(NEO4J_URI, auth=(NEO4J_USER, NEO4J_PASSWORD))
    try:
        async with driver.session(database=NEO4J_DATABASE) as s:
            result = await s.run(
                "MATCH (d:Document) "
                "RETURN d.document_id AS doc_id, d.title AS title, "
                "       d.created_at AS created_at, d.access_level AS access_level "
                "ORDER BY d.created_at DESC LIMIT 100"
            )
            records = await result.data()
        doc_count   = len(records)
        recent_docs = [
            {"doc_id": r["doc_id"], "title": r["title"] or r["doc_id"],
             "created_at": r["created_at"], "access_level": r["access_level"]}
            for r in records[:5]
        ]
        for col in collections:
            if col["name"] == QDRANT_COLLECTION:
                col["doc_count"]   = doc_count
                col["recent_docs"] = recent_docs
    except Exception as exc:
        logger.warning("Neo4j doc count failed: %s", exc)
    finally:
        await driver.close()

    return {"collections": collections, "embedding_model": EMBEDDING_MODEL}


@app.get("/api/collections/{name}/docs")
async def api_collection_docs(name: str):
    """List all Document nodes in Neo4j for a given collection."""
    driver = AsyncGraphDatabase.driver(NEO4J_URI, auth=(NEO4J_USER, NEO4J_PASSWORD))
    try:
        async with driver.session(database=NEO4J_DATABASE) as s:
            result = await s.run(
                "MATCH (d:Document) "
                "OPTIONAL MATCH (d)-[:CONTAINS]->(c:Chunk) "
                "RETURN d.document_id AS doc_id, d.title AS title, "
                "       d.created_at AS created_at, d.access_level AS access_level, "
                "       count(c) AS chunk_count "
                "ORDER BY d.created_at DESC LIMIT 200"
            )
            records = await result.data()
        docs = [
            {"doc_id": r["doc_id"], "title": r["title"] or r["doc_id"],
             "created_at": r["created_at"], "access_level": r["access_level"],
             "chunk_count": r["chunk_count"]}
            for r in records
        ]
        return {"collection": name, "docs": docs}
    except Exception as exc:
        raise HTTPException(status_code=502, detail=str(exc))
    finally:
        await driver.close()


_DOC_ARCHIVE_COLLECTION = "_doc_archive"


@app.delete("/api/collections/{name}/docs/{doc_id}")
async def api_archive_document(name: str, doc_id: str):
    """Archive a single document: move its Qdrant points to _doc_archive, mark Neo4j nodes."""
    now = _now_iso()
    points: list = []

    async with httpx.AsyncClient(timeout=60) as c:
        chk = await c.get(f"{QDRANT_URL}/collections/{name}")
        if chk.status_code == 404:
            raise HTTPException(status_code=404, detail=f"Collection '{name}' not found")

        vec_params = (chk.json().get("result", {})
                      .get("config", {}).get("params", {}).get("vectors", {}))
        vec_size = vec_params.get("size") if isinstance(vec_params, dict) else None
        vec_size = vec_size or await _get_embed_dim()
        distance = vec_params.get("distance", "Cosine") if isinstance(vec_params, dict) else "Cosine"

        # Ensure archive collection exists
        arc_chk = await c.get(f"{QDRANT_URL}/collections/{_DOC_ARCHIVE_COLLECTION}")
        if arc_chk.status_code == 404:
            cr = await c.put(
                f"{QDRANT_URL}/collections/{_DOC_ARCHIVE_COLLECTION}",
                json={"vectors": {"size": vec_size, "distance": distance}},
            )
            cr.raise_for_status()

        # Scroll all points for this document
        offset = None
        while True:
            scroll_body: dict = {
                "limit": 100, "with_payload": True, "with_vector": True,
                "filter": {"must": [{"key": "document_id", "match": {"value": doc_id}}]},
            }
            if offset is not None:
                scroll_body["offset"] = offset
            sr = await c.post(f"{QDRANT_URL}/collections/{name}/points/scroll", json=scroll_body)
            sr.raise_for_status()
            result = sr.json().get("result", {})
            pts = result.get("points", [])
            points.extend(pts)
            offset = result.get("next_page_offset")
            if offset is None:
                break

        if not points:
            raise HTTPException(status_code=404, detail=f"Document '{doc_id}' not found in '{name}'")

        # Stamp archive metadata into each point's payload
        for p in points:
            if p.get("payload") is None:
                p["payload"] = {}
            p["payload"]["_archived_from_collection"] = name
            p["payload"]["_archived_at"] = now

        # Copy to _doc_archive
        ur = await c.put(
            f"{QDRANT_URL}/collections/{_DOC_ARCHIVE_COLLECTION}/points",
            params={"wait": "true"}, json={"points": points},
        )
        ur.raise_for_status()

        # Remove from source collection
        dr = await c.post(
            f"{QDRANT_URL}/collections/{name}/points/delete",
            params={"wait": "true"},
            json={"filter": {"must": [{"key": "document_id", "match": {"value": doc_id}}]}},
        )
        dr.raise_for_status()

    # Neo4j: mark Document as archived and record ArchivedDocument node
    title = doc_id
    driver = AsyncGraphDatabase.driver(NEO4J_URI, auth=(NEO4J_USER, NEO4J_PASSWORD))
    try:
        async with driver.session(database=NEO4J_DATABASE) as s:
            res = await s.run(
                "MATCH (d:Document {document_id: $did}) "
                "SET d.archived = true, d.archived_at = $now, d.archived_from = $col "
                "RETURN d.title AS title",
                did=doc_id, now=now, col=name,
            )
            rec = await res.single()
            if rec and rec["title"]:
                title = rec["title"]
            await s.run(
                "MERGE (a:ArchivedDocument {document_id: $did}) "
                "SET a.title = $title, a.source_collection = $col, "
                "    a.archived_at = $now, a.chunk_count = $cc",
                did=doc_id, title=title, col=name, now=now, cc=len(points),
            )
    except Exception as exc:
        logger.warning("Neo4j archive doc metadata failed: %s", exc)
    finally:
        await driver.close()

    return {
        "status":            "archived",
        "doc_id":            doc_id,
        "title":             title,
        "source_collection": name,
        "chunks_archived":   len(points),
    }


@app.get("/api/doc-archives")
async def api_list_doc_archives():
    """List all individually archived documents."""
    driver = AsyncGraphDatabase.driver(NEO4J_URI, auth=(NEO4J_USER, NEO4J_PASSWORD))
    try:
        async with driver.session(database=NEO4J_DATABASE) as s:
            result = await s.run(
                "MATCH (a:ArchivedDocument) "
                "RETURN a.document_id AS doc_id, a.title AS title, "
                "       a.source_collection AS source_collection, "
                "       a.archived_at AS archived_at, a.chunk_count AS chunk_count "
                "ORDER BY a.archived_at DESC"
            )
            records = await result.data()
    except Exception as exc:
        raise HTTPException(status_code=502, detail=str(exc))
    finally:
        await driver.close()
    return {"doc_archives": [dict(r) for r in records]}


@app.post("/api/doc-archives/{doc_id}/restore")
async def api_restore_doc_archive(doc_id: str, request: Request):
    """Restore an archived document back into its original (or specified) collection."""
    body = await request.json()
    target_collection = (body.get("collection") or "").strip()

    driver = AsyncGraphDatabase.driver(NEO4J_URI, auth=(NEO4J_USER, NEO4J_PASSWORD))
    try:
        async with driver.session(database=NEO4J_DATABASE) as s:
            res = await s.run(
                "MATCH (a:ArchivedDocument {document_id: $did}) "
                "RETURN a.source_collection AS source_collection",
                did=doc_id,
            )
            rec = await res.single()
    except Exception as exc:
        raise HTTPException(status_code=502, detail=str(exc))
    finally:
        await driver.close()

    if not rec:
        raise HTTPException(status_code=404, detail=f"Archived document '{doc_id}' not found")

    restore_to = target_collection or rec["source_collection"]
    points: list = []

    async with httpx.AsyncClient(timeout=60) as c:
        arc_chk = await c.get(f"{QDRANT_URL}/collections/{_DOC_ARCHIVE_COLLECTION}")
        if arc_chk.status_code == 404:
            raise HTTPException(status_code=404, detail="Document archive not found in Qdrant")

        # Ensure target collection exists
        tgt_chk = await c.get(f"{QDRANT_URL}/collections/{restore_to}")
        if tgt_chk.status_code == 404:
            vec_params = (arc_chk.json().get("result", {})
                          .get("config", {}).get("params", {}).get("vectors", {}))
            vec_size = vec_params.get("size") if isinstance(vec_params, dict) else None
            vec_size = vec_size or await _get_embed_dim()
            distance = vec_params.get("distance", "Cosine") if isinstance(vec_params, dict) else "Cosine"
            cr = await c.put(
                f"{QDRANT_URL}/collections/{restore_to}",
                json={"vectors": {"size": vec_size, "distance": distance}},
            )
            cr.raise_for_status()

        # Scroll archived points for this document
        offset = None
        while True:
            scroll_body: dict = {
                "limit": 100, "with_payload": True, "with_vector": True,
                "filter": {"must": [{"key": "document_id", "match": {"value": doc_id}}]},
            }
            if offset is not None:
                scroll_body["offset"] = offset
            sr = await c.post(
                f"{QDRANT_URL}/collections/{_DOC_ARCHIVE_COLLECTION}/points/scroll",
                json=scroll_body,
            )
            sr.raise_for_status()
            result = sr.json().get("result", {})
            pts = result.get("points", [])
            points.extend(pts)
            offset = result.get("next_page_offset")
            if offset is None:
                break

        if not points:
            raise HTTPException(status_code=404, detail=f"No archived vectors found for '{doc_id}'")

        # Strip archive metadata before restoring
        for p in points:
            if p.get("payload"):
                p["payload"].pop("_archived_from_collection", None)
                p["payload"].pop("_archived_at", None)

        # Copy to target collection
        ur = await c.put(
            f"{QDRANT_URL}/collections/{restore_to}/points",
            params={"wait": "true"}, json={"points": points},
        )
        ur.raise_for_status()

        # Delete from archive
        dr = await c.post(
            f"{QDRANT_URL}/collections/{_DOC_ARCHIVE_COLLECTION}/points/delete",
            params={"wait": "true"},
            json={"filter": {"must": [{"key": "document_id", "match": {"value": doc_id}}]}},
        )
        dr.raise_for_status()

    # Neo4j cleanup
    driver = AsyncGraphDatabase.driver(NEO4J_URI, auth=(NEO4J_USER, NEO4J_PASSWORD))
    try:
        async with driver.session(database=NEO4J_DATABASE) as s:
            await s.run(
                "MATCH (d:Document {document_id: $did}) "
                "REMOVE d.archived REMOVE d.archived_at REMOVE d.archived_from",
                did=doc_id,
            )
            await s.run(
                "MATCH (a:ArchivedDocument {document_id: $did}) DETACH DELETE a",
                did=doc_id,
            )
    except Exception as exc:
        logger.warning("Neo4j restore doc cleanup failed: %s", exc)
    finally:
        await driver.close()

    return {"status": "restored", "doc_id": doc_id, "collection": restore_to, "points_restored": len(points)}


@app.post("/api/collections")
async def api_create_collection(request: Request):
    """Create a new empty Qdrant collection."""
    body = await request.json()
    name = (body.get("name") or "").strip()
    if not name or "/" in name or ".." in name:
        raise HTTPException(status_code=400, detail="Invalid collection name")
    vec_size = int(body.get("vector_size", 0)) or await _get_embed_dim()
    distance = body.get("distance", "Cosine")
    async with httpx.AsyncClient(timeout=30) as c:
        chk = await c.get(f"{QDRANT_URL}/collections/{name}")
        if chk.status_code == 200:
            raise HTTPException(status_code=409, detail=f"Collection '{name}' already exists")
        resp = await c.put(
            f"{QDRANT_URL}/collections/{name}",
            json={"vectors": {"size": vec_size, "distance": distance}},
        )
        resp.raise_for_status()
    METRIC_COLLECTION_OPS.labels(operation="create").inc()
    return {"status": "ok", "name": name}


@app.patch("/api/collections/{name}")
async def api_rename_collection(name: str, request: Request):
    """Rename a Qdrant collection: create new, scroll-copy all points, delete old."""
    body = await request.json()
    new_name = (body.get("name") or "").strip()
    if not new_name or "/" in new_name or ".." in new_name:
        raise HTTPException(status_code=400, detail="Invalid collection name")
    if new_name == name:
        return {"status": "ok", "name": new_name, "points_copied": 0}

    async with httpx.AsyncClient(timeout=120) as c:
        info_resp = await c.get(f"{QDRANT_URL}/collections/{name}")
        if info_resp.status_code == 404:
            raise HTTPException(status_code=404, detail=f"Collection '{name}' not found")
        vec_params = (info_resp.json().get("result", {})
                      .get("config", {}).get("params", {}).get("vectors", {}))
        vec_size = vec_params.get("size") if isinstance(vec_params, dict) else None
        vec_size = vec_size or await _get_embed_dim()
        distance = vec_params.get("distance", "Cosine") if isinstance(vec_params, dict) else "Cosine"

        chk = await c.get(f"{QDRANT_URL}/collections/{new_name}")
        if chk.status_code == 200:
            raise HTTPException(status_code=409, detail=f"Collection '{new_name}' already exists")

        cr = await c.put(
            f"{QDRANT_URL}/collections/{new_name}",
            json={"vectors": {"size": vec_size, "distance": distance}},
        )
        cr.raise_for_status()

        offset = None
        total = 0
        while True:
            scroll_body: dict = {"limit": 100, "with_payload": True, "with_vector": True}
            if offset is not None:
                scroll_body["offset"] = offset
            sr = await c.post(
                f"{QDRANT_URL}/collections/{name}/points/scroll", json=scroll_body
            )
            sr.raise_for_status()
            result = sr.json().get("result", {})
            pts = result.get("points", [])
            if pts:
                ur = await c.put(
                    f"{QDRANT_URL}/collections/{new_name}/points",
                    params={"wait": "true"},
                    json={"points": pts},
                )
                ur.raise_for_status()
                total += len(pts)
            offset = result.get("next_page_offset")
            if offset is None:
                break

        await c.delete(f"{QDRANT_URL}/collections/{name}")

    METRIC_COLLECTION_OPS.labels(operation="rename").inc()
    return {"status": "ok", "name": new_name, "points_copied": total}


@app.delete("/api/collections/{name}")
async def api_archive_collection(name: str):
    """Archive a collection: copy to _archived_* in Qdrant, mark Neo4j docs, record in graph."""
    ts_ms = int(datetime.now(timezone.utc).timestamp() * 1000)
    archive_qname = f"_archived_{name}_{ts_ms}"
    now = _now_iso()
    doc_ids: set = set()
    points_count = 0

    async with httpx.AsyncClient(timeout=120) as c:
        info_resp = await c.get(f"{QDRANT_URL}/collections/{name}")
        if info_resp.status_code == 404:
            raise HTTPException(status_code=404, detail=f"Collection '{name}' not found")
        vec_params = (info_resp.json().get("result", {})
                      .get("config", {}).get("params", {}).get("vectors", {}))
        vec_size = vec_params.get("size") if isinstance(vec_params, dict) else None
        vec_size = vec_size or await _get_embed_dim()
        distance = vec_params.get("distance", "Cosine") if isinstance(vec_params, dict) else "Cosine"

        cr = await c.put(
            f"{QDRANT_URL}/collections/{archive_qname}",
            json={"vectors": {"size": vec_size, "distance": distance}},
        )
        cr.raise_for_status()

        offset = None
        while True:
            scroll_body: dict = {"limit": 100, "with_payload": True, "with_vector": True}
            if offset is not None:
                scroll_body["offset"] = offset
            sr = await c.post(
                f"{QDRANT_URL}/collections/{name}/points/scroll", json=scroll_body
            )
            sr.raise_for_status()
            result = sr.json().get("result", {})
            pts = result.get("points", [])
            if pts:
                for p in pts:
                    did = (p.get("payload") or {}).get("document_id")
                    if did:
                        doc_ids.add(did)
                ur = await c.put(
                    f"{QDRANT_URL}/collections/{archive_qname}/points",
                    params={"wait": "true"}, json={"points": pts},
                )
                ur.raise_for_status()
                points_count += len(pts)
            offset = result.get("next_page_offset")
            if offset is None:
                break

        await c.delete(f"{QDRANT_URL}/collections/{name}")

    doc_ids_list = list(doc_ids)
    driver = AsyncGraphDatabase.driver(NEO4J_URI, auth=(NEO4J_USER, NEO4J_PASSWORD))
    try:
        async with driver.session(database=NEO4J_DATABASE) as s:
            if doc_ids_list:
                await s.run(
                    "UNWIND $ids AS did "
                    "MATCH (d:Document {document_id: did}) "
                    "SET d.archived = true, d.archived_at = $now, d.archived_from = $orig",
                    ids=doc_ids_list, now=now, orig=name,
                )
            await s.run(
                "CREATE (:ArchivedCollection {"
                "  qdrant_name: $qname, original_name: $oname,"
                "  archived_at: $now, doc_ids: $doc_ids, points_count: $pc"
                "})",
                qname=archive_qname, oname=name, now=now,
                doc_ids=doc_ids_list, pc=points_count,
            )
    except Exception as exc:
        logger.warning("Neo4j archive metadata failed: %s", exc)
    finally:
        await driver.close()

    METRIC_COLLECTION_OPS.labels(operation="archive").inc()
    return {
        "status":        "archived",
        "archive_name":  archive_qname,
        "original_name": name,
        "points":        points_count,
        "documents":     len(doc_ids_list),
    }


@app.get("/api/archives")
async def api_list_archives():
    """List all archived collections recorded in Neo4j."""
    driver = AsyncGraphDatabase.driver(NEO4J_URI, auth=(NEO4J_USER, NEO4J_PASSWORD))
    try:
        async with driver.session(database=NEO4J_DATABASE) as s:
            result = await s.run(
                "MATCH (a:ArchivedCollection) "
                "RETURN a.qdrant_name AS qdrant_name, a.original_name AS original_name, "
                "       a.archived_at AS archived_at, a.doc_ids AS doc_ids, "
                "       a.points_count AS points_count "
                "ORDER BY a.archived_at DESC"
            )
            records = await result.data()
    except Exception as exc:
        raise HTTPException(status_code=502, detail=str(exc))
    finally:
        await driver.close()

    archives = []
    async with httpx.AsyncClient(timeout=10) as c:
        for r in records:
            chk = await c.get(f"{QDRANT_URL}/collections/{r['qdrant_name']}")
            archives.append({
                "qdrant_name":   r["qdrant_name"],
                "original_name": r["original_name"],
                "archived_at":   r["archived_at"],
                "points_count":  r["points_count"] or 0,
                "doc_count":     len(r["doc_ids"] or []),
                "qdrant_exists": chk.status_code == 200,
            })
    return {"archives": archives}


@app.post("/api/archives/{archive_name}/restore")
async def api_restore_archive(archive_name: str, request: Request):
    """Restore an archived collection to active, with optional rename."""
    body = await request.json()
    restore_name = (body.get("name") or "").strip()

    driver = AsyncGraphDatabase.driver(NEO4J_URI, auth=(NEO4J_USER, NEO4J_PASSWORD))
    try:
        async with driver.session(database=NEO4J_DATABASE) as s:
            res = await s.run(
                "MATCH (a:ArchivedCollection {qdrant_name: $qname}) "
                "RETURN a.original_name AS original_name, a.doc_ids AS doc_ids",
                qname=archive_name,
            )
            rec = await res.single()
    except Exception as exc:
        raise HTTPException(status_code=502, detail=str(exc))
    finally:
        await driver.close()

    if not rec:
        raise HTTPException(status_code=404, detail=f"Archive '{archive_name}' not found")

    original_name = rec["original_name"]
    doc_ids = list(rec["doc_ids"] or [])
    target_name = restore_name or original_name

    async with httpx.AsyncClient(timeout=120) as c:
        info_resp = await c.get(f"{QDRANT_URL}/collections/{archive_name}")
        if info_resp.status_code == 404:
            raise HTTPException(status_code=404, detail=f"Qdrant archive '{archive_name}' not found")
        vec_params = (info_resp.json().get("result", {})
                      .get("config", {}).get("params", {}).get("vectors", {}))
        vec_size = vec_params.get("size") if isinstance(vec_params, dict) else None
        vec_size = vec_size or await _get_embed_dim()
        distance = vec_params.get("distance", "Cosine") if isinstance(vec_params, dict) else "Cosine"

        chk = await c.get(f"{QDRANT_URL}/collections/{target_name}")
        if chk.status_code == 200:
            raise HTTPException(
                status_code=409, detail=f"Collection '{target_name}' already exists"
            )

        cr = await c.put(
            f"{QDRANT_URL}/collections/{target_name}",
            json={"vectors": {"size": vec_size, "distance": distance}},
        )
        cr.raise_for_status()

        offset = None
        total = 0
        while True:
            scroll_body: dict = {"limit": 100, "with_payload": True, "with_vector": True}
            if offset is not None:
                scroll_body["offset"] = offset
            sr = await c.post(
                f"{QDRANT_URL}/collections/{archive_name}/points/scroll", json=scroll_body
            )
            sr.raise_for_status()
            result = sr.json().get("result", {})
            pts = result.get("points", [])
            if pts:
                ur = await c.put(
                    f"{QDRANT_URL}/collections/{target_name}/points",
                    params={"wait": "true"}, json={"points": pts},
                )
                ur.raise_for_status()
                total += len(pts)
            offset = result.get("next_page_offset")
            if offset is None:
                break

        await c.delete(f"{QDRANT_URL}/collections/{archive_name}")

    driver = AsyncGraphDatabase.driver(NEO4J_URI, auth=(NEO4J_USER, NEO4J_PASSWORD))
    try:
        async with driver.session(database=NEO4J_DATABASE) as s:
            if doc_ids:
                await s.run(
                    "UNWIND $ids AS did "
                    "MATCH (d:Document {document_id: did}) "
                    "REMOVE d.archived REMOVE d.archived_at REMOVE d.archived_from",
                    ids=doc_ids,
                )
            await s.run(
                "MATCH (a:ArchivedCollection {qdrant_name: $qname}) DETACH DELETE a",
                qname=archive_name,
            )
    except Exception as exc:
        logger.warning("Neo4j restore cleanup failed: %s", exc)
    finally:
        await driver.close()

    METRIC_COLLECTION_OPS.labels(operation="restore").inc()
    return {"status": "restored", "name": target_name, "points_restored": total}


@app.get("/api/search")
async def api_search(
    collection: str = Query(...),
    q: str = Query(...),
    top_k: int = Query(8),
):
    """Semantic search against a Qdrant collection.

    Validates that the collection's vector dimension matches the active
    embedding model before searching to prevent garbage scores.
    """
    vector = await embed_text(q)
    if not vector:
        raise HTTPException(status_code=503, detail="Embedding unavailable — check OPENAI_API_KEY")
    try:
        async with httpx.AsyncClient(timeout=30) as c:
            # Dimension guard — reject if collection dim != query dim
            info = await c.get(f"{QDRANT_URL}/collections/{collection}")
            if info.status_code == 404:
                return {"results": [], "error": f"Collection '{collection}' not found"}
            col_dim = (info.json().get("result", {})
                       .get("config", {}).get("params", {})
                       .get("vectors", {}).get("size"))
            if col_dim and col_dim != len(vector):
                raise HTTPException(
                    status_code=422,
                    detail=f"Dimension mismatch: collection '{collection}' has {col_dim}-dim "
                           f"vectors but current embedding model produces {len(vector)}-dim. "
                           f"Re-embed the collection or switch EMBEDDING_BACKEND.",
                )
            now = _now_iso()
            temporal_filter = {
                "must": [
                    {"should": [
                        {"is_empty": {"key": "valid_to"}},
                        {"key": "valid_to", "range": {"gte": now}},
                    ]},
                    {"should": [
                        {"is_empty": {"key": "valid_from"}},
                        {"key": "valid_from", "range": {"lte": now}},
                    ]},
                ],
            }
            resp = await c.post(
                f"{QDRANT_URL}/collections/{collection}/points/search",
                json={"vector": vector, "limit": top_k, "with_payload": True,
                      "filter": temporal_filter},
            )
            if resp.status_code == 404:
                return {"results": [], "error": f"Collection '{collection}' not found"}
            resp.raise_for_status()
            hits = resp.json().get("result", [])
    except Exception as exc:
        raise HTTPException(status_code=502, detail=str(exc))

    results = [
        {
            "score":            h.get("score", 0),
            "text":             h.get("payload", {}).get("text", ""),
            "neo4j_element_id": h.get("payload", {}).get("neo4j_element_id", ""),
            "doc_title":        h.get("payload", {}).get("document_title", ""),
            "position":         h.get("payload", {}).get("position"),
        }
        for h in hits
    ]
    return {"collection": collection, "query": q, "results": results}


@app.post("/api/ingest")
async def api_ingest(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    title: str = Form(""),
    agent_id: str = Form(""),
    access_level: str = Form("public"),
    enable_docint: str = Form("false"),   # "true" → run full Document Intelligence
    collection: str = Form(""),           # target Qdrant collection (defaults to QDRANT_COLLECTION)
    valid_hours: float = Form(0.0),       # >0 → chunk expires after this many hours
):
    """JSON-returning ingest endpoint used by the knowledge UI."""
    filename    = file.filename or "upload"
    data        = await file.read()
    use_docint  = enable_docint.lower() in ("true", "1", "yes")

    if not data:
        raise HTTPException(status_code=400, detail="Empty file")

    ext = _resolve_ext(filename, data)
    if ext not in _INGEST_ACCEPT:
        raise HTTPException(status_code=415, detail=f"Unsupported type '{ext}'")

    doc_title  = title.strip() or filename
    doc_id     = hashlib.md5(f"{filename}-{datetime.now().isoformat()}".encode()).hexdigest()[:16]
    image_props: dict = {}
    docint_summary: dict = {}

    try:
        if ext in _MEDICAL_EXTS or _is_image(ext):
            # Medical and image formats always use their dedicated handlers —
            # DocInt pipeline does not understand DICOM/NIfTI/pixel data.
            text, image_props = await _extract_image_text(filename, data, ext)
            text = text.strip()
        elif use_docint:
            # Full Document Intelligence pipeline (PDFs, DOCX, scanned docs)
            result: DocIntelResult = await process_document(filename, data)
            text        = result.to_rich_text().strip()
            image_props = result.image_props
            docint_summary = {
                "document_type":    result.document_type,
                "document_subtype": result.document_subtype,
                "is_scanned":       result.is_scanned,
                "pages":            result.pages,
                "language":         result.language,
                "tables":           len(result.tables),
                "form_fields":      len(result.form_fields),
                "entities": {
                    "dates":         result.entities.dates[:5],
                    "amounts":       result.entities.amounts[:5],
                    "names":         result.entities.names[:5],
                    "organizations": result.entities.organizations[:5],
                },
            }
            # Store tables as separate Neo4j nodes (linked after main ingest)
            _docint_tables_pending[doc_id] = result.tables
        else:
            text = _extract_text(filename, data).strip()
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Extraction failed: {exc}")

    if not text:
        raise HTTPException(status_code=422, detail="No text extracted")
    chunks = _chunk_text(text)
    if not chunks:
        raise HTTPException(status_code=422, detail="No chunks produced")
    target_collection = collection.strip() or QDRANT_COLLECTION

    from datetime import timedelta
    valid_to = (
        (datetime.now(timezone.utc) + timedelta(hours=valid_hours)).isoformat()
        if valid_hours > 0 else None
    )
    logger.info(
        "API ingest '%s': %d chunks (ext=%s, docint=%s, collection=%s, valid_to=%s)",
        doc_title, len(chunks), ext, use_docint, target_collection, valid_to,
    )

    try:
        _, chunk_eids = await _neo4j_ingest(
            doc_id, doc_title, filename, chunks, agent_id, access_level, image_props,
            valid_to=valid_to,
        )
    except Exception as exc:
        METRIC_INGEST_TOTAL.labels(status="neo4j_error").inc()
        raise HTTPException(status_code=502, detail=f"Neo4j: {exc}")

    # Persist DocTable nodes for DocInt results
    if doc_id in _docint_tables_pending:
        try:
            await _neo4j_ingest_tables(doc_id, _docint_tables_pending.pop(doc_id))
        except Exception as e:
            logger.warning("DocInt table Neo4j write failed: %s", e)

    try:
        upserted = await _qdrant_ingest(
            target_collection, chunks, chunk_eids, doc_id, doc_title, agent_id, access_level,
            valid_to=valid_to,
        )
    except Exception as exc:
        METRIC_INGEST_TOTAL.labels(status="qdrant_error").inc()
        raise HTTPException(status_code=502, detail=f"Qdrant: {exc}")

    METRIC_INGEST_TOTAL.labels(status="success").inc()
    METRIC_INGEST_CHUNKS.observe(len(chunks))

    # NER entity extraction runs in the background so ingest response is immediate
    background_tasks.add_task(_neo4j_create_mentions, chunk_eids, chunks)

    result_json = {"status": "ok", "document_id": doc_id, "title": doc_title,
                   "filename": filename, "chunks": len(chunks), "qdrant_points": upserted,
                   "collection": target_collection}
    if valid_to:
        result_json["valid_to"] = valid_to
    if image_props:
        result_json["image_metadata"] = image_props
    if docint_summary:
        result_json["docint"] = docint_summary
    return result_json


# In-memory staging for DocTable nodes created during /api/ingest
_docint_tables_pending: dict[str, list] = {}


@app.post("/api/ner-backfill")
async def ner_backfill(
    limit: int = Query(500, description="Max chunks to process per call"),
):
    """Run NER entity extraction on Chunk nodes that have no MENTIONS edges yet.

    Calls Ollama (OLLAMA_MODEL) per chunk, creates :Entity nodes and MENTIONS
    relationships.  Safe to call repeatedly — only touches unprocessed chunks.
    """
    driver = AsyncGraphDatabase.driver(NEO4J_URI, auth=(NEO4J_USER, NEO4J_PASSWORD))
    try:
        async with driver.session(database=NEO4J_DATABASE) as s:
            result = await s.run(
                """
                MATCH (c:Chunk)
                WHERE NOT (c)-[:MENTIONS]->()
                  AND c.text IS NOT NULL
                RETURN elementId(c) AS eid, c.text AS text
                LIMIT $lim
                """,
                lim=limit,
            )
            rows = await result.data()
    finally:
        await driver.close()

    if not rows:
        return {"status": "ok", "processed": 0, "mentions_created": 0,
                "message": "No unprocessed chunks found"}

    chunk_eids  = [r["eid"]  for r in rows]
    chunk_texts = [r["text"] or "" for r in rows]
    mentions = await _neo4j_create_mentions(chunk_eids, chunk_texts)
    return {
        "status":          "ok",
        "processed":       len(rows),
        "mentions_created": mentions,
    }


@app.post("/api/graph-backfill")
async def graph_backfill():
    """Backfill NEXT chunk links and CO_OCCURS entity edges for all existing data.

    Safe to call repeatedly — uses MERGE so relationships are idempotent.
    """
    driver = AsyncGraphDatabase.driver(NEO4J_URI, auth=(NEO4J_USER, NEO4J_PASSWORD))
    try:
        async with driver.session(database=NEO4J_DATABASE) as s:
            # 1. Create NEXT links between consecutive chunks within each document
            next_result = await s.run(
                """
                MATCH (d:Document)-[:CONTAINS]->(c:Chunk)
                WITH d, c ORDER BY c.position
                WITH d, collect(c) AS ordered
                UNWIND range(0, size(ordered) - 2) AS i
                WITH ordered[i] AS a, ordered[i + 1] AS b
                MERGE (a)-[:NEXT]->(b)
                RETURN count(*) AS next_created
                """
            )
            next_count = (await next_result.single())["next_created"]

            # 2. Create CO_OCCURS edges between entities sharing a chunk
            cooc_result = await s.run(
                """
                MATCH (c:Chunk)-[:MENTIONS]->(e1:Entity)
                MATCH (c)-[:MENTIONS]->(e2:Entity)
                WHERE elementId(e1) < elementId(e2)
                MERGE (e1)-[r:CO_OCCURS]->(e2)
                  ON CREATE SET r.weight = 1, r.created_at = $now
                  ON MATCH  SET r.weight = r.weight + 1
                RETURN count(*) AS cooc_created
                """,
                now=_now_iso(),
            )
            cooc_count = (await cooc_result.single())["cooc_created"]
    finally:
        await driver.close()

    return {
        "status": "ok",
        "next_links": next_count,
        "co_occurs_edges": cooc_count,
    }


@app.post("/api/buddy-memory-migrate")
async def buddy_memory_migrate(
    source: str = Query("buddy_memory", description="Collection to migrate"),
    batch_size: int = Query(32, description="Re-embedding batch size"),
):
    """Migrate a Qdrant collection from 3072-dim (OpenAI large) to 768-dim (nomic-embed-text).

    Steps
    -----
    1. Scroll all points from *source* (preserving full payloads + old vectors).
    2. Back them up as ``{source}_backup_3072`` — old dim intact, safe to restore.
    3. Re-embed every point's ``text`` payload field via the configured Ollama
       embedding model (OLLAMA_EMBEDDING_MODEL, typically nomic-embed-text → 768-dim).
    4. Delete and recreate *source* with the new vector size.
    5. Upsert all points with new vectors.

    Returns a summary including old/new dimensions and backup collection name.
    """
    backup_name = f"{source}_backup_3072"

    # ── 1. Scroll all existing points ─────────────────────────────────────────
    all_points: list[dict] = []
    offset = None
    async with httpx.AsyncClient(timeout=60) as c:
        while True:
            body: dict = {"limit": 256, "with_payload": True, "with_vector": True}
            if offset is not None:
                body["offset"] = offset
            resp = await c.post(
                f"{QDRANT_URL}/collections/{source}/points/scroll",
                json=body,
            )
            if resp.status_code == 404:
                raise HTTPException(status_code=404, detail=f"Collection '{source}' not found")
            resp.raise_for_status()
            data = resp.json()["result"]
            batch = data.get("points", [])
            all_points.extend(batch)
            offset = data.get("next_page_offset")
            if not offset or not batch:
                break

    if not all_points:
        raise HTTPException(status_code=422, detail=f"Collection '{source}' exists but has no points")

    old_dim = len(all_points[0]["vector"])

    # ── 2. Create backup of original vectors ──────────────────────────────────
    async with httpx.AsyncClient(timeout=60) as c:
        await c.delete(f"{QDRANT_URL}/collections/{backup_name}")
        r = await c.put(
            f"{QDRANT_URL}/collections/{backup_name}",
            json={"vectors": {"size": old_dim, "distance": "Cosine"}},
        )
        r.raise_for_status()
        for i in range(0, len(all_points), 256):
            chunk = all_points[i : i + 256]
            r = await c.put(
                f"{QDRANT_URL}/collections/{backup_name}/points",
                params={"wait": "true"},
                json={"points": [
                    {"id": p["id"], "vector": p["vector"], "payload": p.get("payload", {})}
                    for p in chunk
                ]},
            )
            r.raise_for_status()
    logger.info("buddy-memory-migrate: backed up %d points → %s", len(all_points), backup_name)

    # ── 3. Re-embed via Ollama (nomic-embed-text → 768-dim) ───────────────────
    texts = [
        (p.get("payload") or {}).get("text") or (p.get("payload") or {}).get("content") or ""
        for p in all_points
    ]
    new_vectors: list[list[float]] = []
    for i in range(0, len(texts), batch_size):
        vecs = await _embed_batch(texts[i : i + batch_size])
        new_vectors.extend(vecs)

    new_dim = len(new_vectors[0]) if new_vectors else 768

    # ── 4. Recreate source collection with new dimension ──────────────────────
    async with httpx.AsyncClient(timeout=60) as c:
        await c.delete(f"{QDRANT_URL}/collections/{source}")
        r = await c.put(
            f"{QDRANT_URL}/collections/{source}",
            json={"vectors": {"size": new_dim, "distance": "Cosine"}},
        )
        r.raise_for_status()

    # ── 5. Upsert re-embedded points ──────────────────────────────────────────
    async with httpx.AsyncClient(timeout=120) as c:
        for i in range(0, len(all_points), 256):
            chunk   = all_points[i : i + 256]
            ch_vecs = new_vectors[i : i + 256]
            r = await c.put(
                f"{QDRANT_URL}/collections/{source}/points",
                params={"wait": "true"},
                json={"points": [
                    {"id": p["id"], "vector": vec, "payload": p.get("payload", {})}
                    for p, vec in zip(chunk, ch_vecs)
                ]},
            )
            r.raise_for_status()

    logger.info(
        "buddy-memory-migrate: %d points migrated %d→%d dim in '%s'",
        len(all_points), old_dim, new_dim, source,
    )
    return {
        "status":     "ok",
        "collection": source,
        "migrated":   len(all_points),
        "old_dim":    old_dim,
        "new_dim":    new_dim,
        "backup":     backup_name,
    }


async def _neo4j_ingest_tables(doc_id: str, tables: list) -> None:
    """Create DocTable nodes in Neo4j linked to the parent Document."""
    if not tables:
        return
    driver = AsyncGraphDatabase.driver(NEO4J_URI, auth=(NEO4J_USER, NEO4J_PASSWORD))
    now = _now_iso()
    try:
        async with driver.session(database=NEO4J_DATABASE) as s:
            for i, t in enumerate(tables):
                await s.run(
                    """MATCH (d:Document {document_id: $doc_id})
                       CREATE (t:DocTable {
                         table_id:   $tid,
                         caption:    $caption,
                         markdown:   $markdown,
                         row_count:  $rows,
                         col_count:  $cols,
                         page:       $page,
                         created_at: $now
                       })
                       CREATE (d)-[:HAS_TABLE]->(t)""",
                    doc_id=doc_id, tid=f"{doc_id}-table-{i}",
                    caption=t.caption, markdown=t.markdown,
                    rows=t.row_count, cols=t.col_count,
                    page=t.page, now=now,
                )
    finally:
        await driver.close()


@app.post("/api/docint")
async def api_docint(
    file: UploadFile = File(...),
    extract_tables:   str = Form("true"),
    extract_forms:    str = Form("true"),
    extract_entities: str = Form("true"),
    classify:         str = Form("true"),
    vision:           str = Form("true"),
):
    """
    Standalone Document Intelligence endpoint.
    Returns rich structured output without ingesting into Qdrant/Neo4j.
    """
    filename = file.filename or "upload"
    data     = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Empty file")

    ext = _resolve_ext(filename, data)
    if ext not in _INGEST_ACCEPT:
        raise HTTPException(status_code=415, detail=f"Unsupported type '{ext}'")

    def _bool(v: str) -> bool:
        return v.lower() in ("true", "1", "yes")

    try:
        result: DocIntelResult = await process_document(
            filename, data,
            extract_tables=_bool(extract_tables),
            extract_forms=_bool(extract_forms),
            extract_entities=_bool(extract_entities),
            classify=_bool(classify),
            vision=_bool(vision),
        )
    except Exception as exc:
        raise HTTPException(status_code=422, detail=str(exc))

    return {
        "filename":         filename,
        "document_type":    result.document_type,
        "document_subtype": result.document_subtype,
        "type_confidence":  result.type_confidence,
        "language":         result.language,
        "is_scanned":       result.is_scanned,
        "pages":            result.pages,
        "char_count":       result.char_count,
        "tables": [
            {"caption": t.caption, "markdown": t.markdown,
             "rows": t.row_count, "cols": t.col_count, "page": t.page}
            for t in result.tables
        ],
        "form_fields": result.form_fields,
        "entities": {
            "dates":         result.entities.dates,
            "amounts":       result.entities.amounts,
            "names":         result.entities.names,
            "organizations": result.entities.organizations,
            "identifiers":   result.entities.identifiers,
            "locations":     result.entities.locations,
        },
        "text_preview": result.text[:500],
    }


@app.get("/health")
async def health():
    return {
        "status":               "ok",
        "rag":                  ENABLE_RAG,
        "confidence_threshold": CONFIDENCE_THRESHOLD,
        "agent_id":             AGENT_ID or None,
    }


# ── Document Ingest ────────────────────────────────────────────────────────────

_INGEST_ACCEPT = {
    # Documents
    ".txt", ".md", ".pdf", ".docx", ".json", ".csv",
    # Regular images
    ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp", ".tiff", ".tif", ".avif",
    # Medical imaging
    ".dcm", ".dicom", ".nii", ".nii.gz",
}

_IMAGE_EXTS   = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp", ".tiff", ".tif", ".avif"}
_MEDICAL_EXTS = {".dcm", ".dicom", ".nii", ".nii.gz"}


def _resolve_ext(filename: str, data: bytes) -> str:
    """Return the canonical extension, handling .nii.gz and DICOM magic."""
    name = filename.lower()
    if name.endswith(".nii.gz"):
        return ".nii.gz"
    ext = os.path.splitext(name)[1]
    # DICOM files often arrive without extension — detect by preamble magic
    if ext not in _INGEST_ACCEPT and len(data) >= 132 and data[128:132] == b"DICM":
        return ".dcm"
    return ext


def _is_image(ext: str) -> bool:
    return ext in _IMAGE_EXTS or ext in _MEDICAL_EXTS


# ── Vision model helpers ───────────────────────────────────────────────────────

_VISION_MAX_PX     = int(os.getenv("VISION_MAX_PX", "1024"))
_VISION_MAX_TOKENS = int(os.getenv("VISION_MAX_TOKENS", "1200"))
_NO_VISION         = os.getenv("NO_VISION", "0") == "1"


async def _call_vision_api(png_b64: str, prompt: str) -> str:
    """Send a PNG (base64) to the configured vision model and return the description."""
    if _NO_VISION:
        return "[Vision analysis skipped — NO_VISION=1]"

    if LLM_BACKEND == "anthropic" and ANTHROPIC_API_KEY:
        async with httpx.AsyncClient(timeout=60) as c:
            r = await c.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "x-api-key":         ANTHROPIC_API_KEY,
                    "anthropic-version": "2023-06-01",
                    "Content-Type":      "application/json",
                },
                json={
                    "model":      os.getenv("VISION_MODEL", "claude-3-5-sonnet-20241022"),
                    "max_tokens": _VISION_MAX_TOKENS,
                    "messages": [{
                        "role": "user",
                        "content": [
                            {"type": "image",
                             "source": {"type": "base64", "media_type": "image/png",
                                        "data": png_b64}},
                            {"type": "text", "text": prompt},
                        ],
                    }],
                },
            )
            r.raise_for_status()
            return r.json()["content"][0]["text"].strip()

    # OpenAI fallback
    if not OPENAI_API_KEY:
        return "[Vision unavailable — no API key configured]"
    async with httpx.AsyncClient(timeout=60) as c:
        r = await c.post(
            "https://api.openai.com/v1/chat/completions",
            headers={"Authorization": f"Bearer {OPENAI_API_KEY}",
                     "Content-Type": "application/json"},
            json={
                "model":      os.getenv("VISION_MODEL", "gpt-4o"),
                "max_tokens": _VISION_MAX_TOKENS,
                "messages": [{
                    "role": "user",
                    "content": [
                        {"type": "image_url",
                         "image_url": {"url": f"data:image/png;base64,{png_b64}",
                                       "detail": "high"}},
                        {"type": "text", "text": prompt},
                    ],
                }],
            },
        )
        r.raise_for_status()
        return r.json()["choices"][0]["message"]["content"].strip()


def _pil_to_png_b64(img) -> str:
    from PIL import Image  # noqa: F811
    import base64
    img.thumbnail((_VISION_MAX_PX, _VISION_MAX_PX))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return base64.b64encode(buf.getvalue()).decode()


def _normalize_to_uint8(arr):
    import numpy as np
    arr = arr.astype(np.float32)
    mn, mx = arr.min(), arr.max()
    if mx == mn:
        return np.zeros_like(arr, dtype=np.uint8)
    return ((arr - mn) / (mx - mn) * 255).astype(np.uint8)


def _apply_wl(arr, wc: float, ww: float):
    import numpy as np
    low, high = wc - ww / 2, wc + ww / 2
    return np.clip((arr.astype(np.float32) - low) / (ww) * 255, 0, 255).astype(np.uint8)


async def _extract_image_text(filename: str, data: bytes, ext: str) -> tuple[str, dict]:
    """
    Returns (text, image_props) for all image formats.
    text      — descriptive string ready for chunking
    image_props — flat dict stored on the Neo4j Document node as image_* fields
    """
    import base64

    # ── DICOM ──────────────────────────────────────────────────────────────────
    if ext in (".dcm", ".dicom"):
        try:
            import pydicom
            from PIL import Image
            import numpy as np

            ds = pydicom.dcmread(io.BytesIO(data), force=True)

            def _tag(attr, default=""):
                v = getattr(ds, attr, default)
                return str(v).strip() if v else default

            def _flt(attr):
                try:
                    v = getattr(ds, attr, None)
                    if v is None:
                        return None
                    if hasattr(v, "__iter__") and not isinstance(v, str):
                        v = list(v)[0]
                    return float(v)
                except Exception:
                    return None

            modality    = _tag("Modality")
            study_desc  = _tag("StudyDescription")
            series_desc = _tag("SeriesDescription")
            institution = _tag("InstitutionName")
            manufacturer = _tag("Manufacturer")
            protocol    = _tag("ProtocolName")
            rows        = _tag("Rows")
            cols        = _tag("Columns")
            bits        = _tag("BitsStored")
            n_frames    = int(getattr(ds, "NumberOfFrames", 1) or 1)
            photometric = _tag("PhotometricInterpretation")
            slice_thick = _tag("SliceThickness")
            pixel_sp    = _tag("PixelSpacing")
            kvp         = _tag("KVP")

            image_props = {
                "image_format": "DICOM",
                "image_modality": modality,
                "image_study_description": study_desc,
                "image_series_description": series_desc,
                "image_institution": institution,
                "image_manufacturer": manufacturer,
                "image_protocol": protocol,
                "image_rows": rows, "image_cols": cols,
                "image_bits": bits, "image_frames": str(n_frames),
                "image_photometric": photometric,
                "image_slice_thickness_mm": slice_thick,
                "image_pixel_spacing_mm": pixel_sp,
                "image_kvp": kvp,
            }

            meta_lines = [
                f"[MEDICAL IMAGE: DICOM{f' / {modality}' if modality else ''}]",
                "",
                "=== DICOM METADATA ===",
                f"Modality            : {modality}" if modality else None,
                f"Study Description   : {study_desc}" if study_desc else None,
                f"Series Description  : {series_desc}" if series_desc else None,
                f"Institution         : {institution}" if institution else None,
                f"Manufacturer        : {manufacturer}" if manufacturer else None,
                f"Protocol            : {protocol}" if protocol else None,
                f"KVP                 : {kvp} kV" if kvp else None,
                f"Slice Thickness     : {slice_thick} mm" if slice_thick else None,
                f"Pixel Spacing       : {pixel_sp} mm" if pixel_sp else None,
                f"Dimensions          : {cols} × {rows} px" if rows and cols else None,
                f"Frames              : {n_frames}" if n_frames > 1 else None,
                f"Bit Depth           : {bits}-bit" if bits else None,
                f"Photometric         : {photometric}" if photometric else None,
            ]
            meta_text = "\n".join(l for l in meta_lines if l is not None)

            # Pixel rendering
            vision_text = ""
            try:
                px = ds.pixel_array  # (rows, cols) or (frames, rows, cols)
                if px.ndim == 3:
                    px = px[px.shape[0] // 2]  # middle frame

                wc = _flt("WindowCenter")
                ww = _flt("WindowWidth")
                gray8 = _apply_wl(px, wc, ww) if (wc is not None and ww and ww > 0) \
                    else _normalize_to_uint8(px)

                # Handle RGB DICOM
                if gray8.ndim == 3:
                    img = Image.fromarray(gray8)
                else:
                    img = Image.fromarray(gray8, mode="L")

                b64 = _pil_to_png_b64(img)
                prompt = (
                    f"This is a medical DICOM image (modality: {modality or 'unknown'}). "
                    "Describe clinically: (1) anatomical region and structures visible, "
                    "(2) notable findings or abnormalities (describe objectively, do not diagnose), "
                    "(3) image orientation and quality, (4) any visible text, annotations, "
                    "measurements, or overlays."
                )
                vision_text = await _call_vision_api(b64, prompt)
            except Exception as e:
                vision_text = f"[Pixel rendering failed: {e}]"

            text = meta_text
            if vision_text:
                text += f"\n\n=== VISUAL ANALYSIS ===\n{vision_text}"
            return text, {k: v for k, v in image_props.items() if v}

        except ImportError as e:
            raise RuntimeError(
                f"Missing dependency for DICOM: {e}. "
                "Run: pip install pydicom Pillow numpy"
            )

    # ── NIfTI ──────────────────────────────────────────────────────────────────
    if ext in (".nii", ".nii.gz"):
        try:
            import nibabel as nib
            import numpy as np
            import tempfile
            from PIL import Image

            suffix = ".nii.gz" if ext == ".nii.gz" else ".nii"
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
                tmp.write(data)
                tmp_path = tmp.name
            try:
                nii = nib.load(tmp_path)
                hdr = nii.header
                shape = nii.shape  # (nx, ny, nz[, nt])
                vox   = hdr.get_zooms()
                dtype = str(hdr.get_data_dtype())
                descr = hdr.get("descrip", b"").tobytes().decode("utf-8", errors="replace").strip("\x00").strip()
            finally:
                os.unlink(tmp_path)

            nx, ny, nz = (shape + (1, 1))[:3]
            nt = shape[3] if len(shape) > 3 else 1
            dx, dy, dz = (vox + (0, 0, 0))[:3]

            image_props = {
                "image_format": "NIfTI",
                "image_dims": f"{nx}x{ny}x{nz}" + (f"x{nt}" if nt > 1 else ""),
                "image_voxel_size_mm": f"{dx:.3f}x{dy:.3f}x{dz:.3f}",
                "image_dtype": dtype,
                "image_description": descr,
            }

            meta_text = "\n".join(filter(None, [
                "[MEDICAL IMAGE: NIfTI]",
                "",
                "=== NIfTI HEADER ===",
                f"Dimensions   : {nx} × {ny} × {nz}" + (f" × {nt} (time)" if nt > 1 else ""),
                f"Voxel Size   : {dx:.3f} × {dy:.3f} × {dz:.3f} mm",
                f"Data Type    : {dtype}",
                f"Description  : {descr}" if descr else None,
            ]))

            vision_text = ""
            try:
                arr = nii.get_fdata()
                if arr.ndim >= 3:
                    sl = arr[:, :, arr.shape[2] // 2]
                elif arr.ndim == 2:
                    sl = arr
                else:
                    raise ValueError("Unexpected NIfTI array shape")
                gray8 = _normalize_to_uint8(sl)
                img = Image.fromarray(gray8, mode="L")
                b64 = _pil_to_png_b64(img)
                prompt = (
                    "This is a medical neuroimaging image (NIfTI format). "
                    "Describe: (1) visible brain structures or anatomy, (2) any notable "
                    "features, signal intensities, or abnormalities, (3) image plane and "
                    "orientation, (4) image quality and contrast."
                )
                vision_text = await _call_vision_api(b64, prompt)
            except Exception as e:
                vision_text = f"[Slice rendering failed: {e}]"

            text = meta_text
            if vision_text:
                text += f"\n\n=== VISUAL ANALYSIS ===\n{vision_text}"
            return text, {k: v for k, v in image_props.items() if v}

        except ImportError as e:
            raise RuntimeError(
                f"Missing dependency for NIfTI: {e}. "
                "Run: pip install nibabel numpy Pillow"
            )

    # ── Regular images ──────────────────────────────────────────────────────────
    try:
        from PIL import Image
        import base64

        img = Image.open(io.BytesIO(data))
        width, height = img.size
        mode  = img.mode
        fmt   = img.format or ext.lstrip(".")

        image_props = {
            "image_format": fmt.upper(),
            "image_width":  str(width),
            "image_height": str(height),
            "image_mode":   mode,
        }

        # Convert to RGB PNG for vision
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        b64 = _pil_to_png_b64(img)

        prompt = (
            "Describe this image in detail: (1) main subject and overall content, "
            "(2) any text, labels, annotations, or measurements visible, "
            "(3) key visual elements, colors, or patterns, "
            "(4) context or apparent purpose."
        )
        vision_text = await _call_vision_api(b64, prompt)

        text = "\n".join([
            f"[IMAGE: {fmt.upper()}  {width}×{height}px]",
            "",
            vision_text or "[Vision analysis skipped]",
        ])
        return text, image_props

    except ImportError as e:
        raise RuntimeError(f"Missing dependency for images: {e}. Run: pip install Pillow")

_UPLOAD_FORM = """<!doctype html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width,initial-scale=1">
  <title>EedgeAI — Ingest Document</title>
  <style>
    body{{font-family:system-ui,sans-serif;max-width:640px;margin:60px auto;padding:0 24px;background:#0f172a;color:#e2e8f0}}
    h1{{color:#38bdf8;margin-bottom:4px}}
    p{{color:#94a3b8;margin-top:0}}
    form{{background:#1e293b;padding:28px;border-radius:12px;margin-top:24px}}
    label{{display:block;font-size:.85rem;color:#94a3b8;margin-bottom:4px}}
    input,select{{width:100%;padding:8px 12px;border:1px solid #334155;border-radius:6px;
      background:#0f172a;color:#e2e8f0;font-size:.9rem;box-sizing:border-box;margin-bottom:16px}}
    button{{background:#0ea5e9;color:#fff;border:none;padding:10px 28px;
      border-radius:6px;font-size:1rem;cursor:pointer;width:100%}}
    button:hover{{background:#38bdf8}}
    .hint{{font-size:.78rem;color:#64748b;margin-top:-12px;margin-bottom:16px}}
    .accepted{{color:#64748b;font-size:.8rem;margin-top:16px}}
  </style>
</head>
<body>
  <h1>Ingest Document</h1>
  <p>Upload a file to index it into Neo4j + Qdrant for GCOR retrieval.</p>
  <form method="POST" enctype="multipart/form-data">
    <label>File</label>
    <input type="file" name="file" accept=".txt,.md,.pdf,.docx,.json,.csv,.jpg,.jpeg,.png,.gif,.bmp,.webp,.tiff,.tif,.avif,.dcm,.dicom,.nii" required>
    <p class="accepted">Accepted: .txt .md .pdf .docx .json .csv · images: .jpg .png .webp .tiff · medical: .dcm .dicom .nii .nii.gz</p>
    <label>Title (optional — defaults to filename)</label>
    <input type="text" name="title" placeholder="My Document">
    <label>Agent ID (optional — leave blank for shared knowledge)</label>
    <input type="text" name="agent_id" placeholder="">
    <label>Access level</label>
    <select name="access_level">
      <option value="public">public</option>
      <option value="restricted">restricted</option>
    </select>
    <button type="submit">Ingest</button>
  </form>
</body>
</html>"""

_RESULT_TMPL = """<!doctype html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <title>Ingest result</title>
  <style>
    body{{font-family:system-ui,sans-serif;max-width:640px;margin:60px auto;padding:0 24px;background:#0f172a;color:#e2e8f0}}
    h1{{color:{color}}}
    pre{{background:#1e293b;padding:20px;border-radius:8px;overflow-x:auto;font-size:.85rem}}
    a{{color:#38bdf8}}
  </style>
</head>
<body>
  <h1>{heading}</h1>
  <pre>{body}</pre>
  <p><a href="/ingest">&larr; Ingest another</a></p>
</body>
</html>"""


def _extract_text(filename: str, data: bytes) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        return "\n".join(p.extract_text() or "" for p in reader.pages)
    if ext == ".docx":
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)
    if ext == ".json":
        try:
            return json.dumps(json.loads(data), indent=2)
        except Exception:
            pass
    return data.decode("utf-8", errors="replace")


import re as _re

_ENTITY_PATTERNS = [
    r'\b[A-Z][a-z]+ [A-Z][a-z]+(?:\s[A-Z][a-z]+)?\b',  # Person/place names (Title Case)
    r'\b[A-Z]{2,}\b',                                      # Acronyms (e.g. AIOps, GCOR, AWS)
    r'\b\d{4}-\d{2}-\d{2}\b',                             # ISO dates (2026-03-25)
]


def _find_entity_boundaries(text: str) -> set:
    """Return character positions that are mid-entity — avoid splitting here."""
    bad: set = set()
    for pattern in _ENTITY_PATTERNS:
        for m in _re.finditer(pattern, text):
            bad.update(range(m.start(), m.end()))
    return bad


def _chunk_text(text: str, size: int = 2000, overlap: int = 200) -> list[str]:
    bad_positions = _find_entity_boundaries(text)
    chunks, start = [], 0
    while start < len(text):
        end = min(start + size, len(text))
        if end < len(text):
            for sep in ("\n\n", "\n", " "):
                pos = text.rfind(sep, start + size // 2, end)
                if pos != -1 and pos not in bad_positions:
                    end = pos
                    break
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start = end - overlap
        if start >= len(text) - overlap:
            break
    return chunks


# ── NER Entity Extraction ─────────────────────────────────────────────────────

async def _ner_extract(text: str) -> list[dict]:
    """Call Ollama (OLLAMA_MODEL) to extract named entities from a chunk of text.

    Uses response_format=json_object so Ollama is constrained to valid JSON.
    Returns a list of dicts like {"text": "...", "type": "PERSON|ORG|..."}.
    Falls back to [] on any error so ingest is never blocked.
    """
    system = (
        "You are a named entity extractor. "
        'Return a JSON object with a single key "entities" whose value is an array. '
        'Each array item must have "text" (string) and "type" (one of: '
        "PERSON, ORG, LOCATION, DATE, CONCEPT, TECH, PRODUCT, EVENT). "
        "Include at most 15 of the most significant entities. "
        "If there are none, return {\"entities\": []}."
    )
    user = f"Extract entities from:\n\n{text[:3000]}"
    try:
        async with httpx.AsyncClient(timeout=45) as c:
            resp = await c.post(
                f"{OLLAMA_BASE_URL}/v1/chat/completions",
                json={
                    "model":           OLLAMA_MODEL,
                    "messages":        [
                        {"role": "system", "content": system},
                        {"role": "user",   "content": user},
                    ],
                    "temperature":     0,
                    "response_format": {"type": "json_object"},
                },
            )
            resp.raise_for_status()
        content = resp.json()["choices"][0]["message"]["content"].strip()
        parsed = json.loads(content)
        # Accept {"entities": [...]} or a bare array
        if isinstance(parsed, list):
            entities = parsed
        else:
            entities = parsed.get("entities") or parsed.get("Entities") or []
            if not isinstance(entities, list):
                entities = []
        return [
            e for e in entities
            if isinstance(e, dict) and e.get("text") and e.get("type")
        ]
    except Exception as exc:
        logger.warning("NER extraction failed: %s", exc)
        return []


_VALID_ENTITY_TYPES = frozenset(
    {"PERSON", "ORG", "LOCATION", "DATE", "CONCEPT", "TECH", "PRODUCT", "EVENT"}
)


async def _neo4j_create_mentions(chunk_eids: list[str], chunks: list[str]) -> int:
    """Extract entities per chunk and write :Entity nodes + MENTIONS rels to Neo4j.

    Also creates :CO_OCCURS relationships between entities that appear in the
    same chunk, which is the foundation for relational / graph reasoning.

    Returns total number of MENTIONS relationships created.
    Safe to call in a background task — errors are logged, not raised.
    """
    total = 0
    driver = AsyncGraphDatabase.driver(NEO4J_URI, auth=(NEO4J_USER, NEO4J_PASSWORD))
    try:
        async with driver.session(database=NEO4J_DATABASE) as s:
            for eid, text in zip(chunk_eids, chunks):
                try:
                    entities = await _ner_extract(text)
                except Exception as exc:
                    logger.warning("NER failed for chunk %s: %s", eid, exc)
                    continue

                # Deduplicate and validate entities for this chunk
                seen_names: list[str] = []
                for ent in entities:
                    name  = str(ent.get("text", "")).strip()
                    etype = str(ent.get("type", "CONCEPT")).strip().upper()
                    if not name:
                        continue
                    if etype not in _VALID_ENTITY_TYPES:
                        etype = "CONCEPT"
                    await s.run(
                        """
                        MATCH (c:Chunk) WHERE elementId(c) = $eid
                        MERGE (e:Entity {name: $name, type: $etype})
                          ON CREATE SET e.created_at = $now
                        MERGE (c)-[:MENTIONS]->(e)
                        """,
                        eid=eid, name=name, etype=etype, now=_now_iso(),
                    )
                    seen_names.append(name)
                    total += 1

                # Create CO_OCCURS edges between all entity pairs in this chunk
                if len(seen_names) >= 2:
                    await s.run(
                        """
                        MATCH (c:Chunk) WHERE elementId(c) = $eid
                        MATCH (c)-[:MENTIONS]->(e1:Entity)
                        MATCH (c)-[:MENTIONS]->(e2:Entity)
                        WHERE elementId(e1) < elementId(e2)
                        MERGE (e1)-[r:CO_OCCURS]->(e2)
                          ON CREATE SET r.weight = 1, r.created_at = $now
                          ON MATCH  SET r.weight = r.weight + 1
                        """,
                        eid=eid, now=_now_iso(),
                    )
    except Exception as exc:
        logger.error("_neo4j_create_mentions failed: %s", exc)
    finally:
        await driver.close()
    logger.info("NER backfill: %d MENTIONS created for %d chunks", total, len(chunk_eids))
    return total


async def _neo4j_ingest(doc_id: str, title: str, source: str,
                        chunks: list[str], agent_id: str, access_level: str,
                        image_props: dict | None = None,
                        valid_to: str | None = None):
    driver = AsyncGraphDatabase.driver(NEO4J_URI, auth=(NEO4J_USER, NEO4J_PASSWORD))
    now = _now_iso()
    try:
        async with driver.session(database=NEO4J_DATABASE) as s:
            doc_rec = await s.run(
                """CREATE (d:Document {
                     document_id: $doc_id, title: $title, source: $source,
                     agent_id: $agent_id, access_level: $access_level,
                     chunk_count: $cc, created_at: $now
                   })
                   SET d += $image_props
                   RETURN elementId(d) AS eid""",
                doc_id=doc_id, title=title, source=source,
                agent_id=agent_id, access_level=access_level,
                cc=len(chunks), now=now,
                image_props=image_props or {},
            )
            doc_eid = (await doc_rec.single())["eid"]

            chunk_eids = []
            for i, text in enumerate(chunks):
                cid = f"{doc_id}-chunk-{i}"
                c_rec = await s.run(
                    """MATCH (d:Document {document_id: $doc_id})
                       CREATE (c:Chunk {
                         chunk_id: $cid, text: $text, position: $pos,
                         document_id: $doc_id, document_title: $title,
                         agent_id: $agent_id, access_level: $access_level,
                         confidence: 1.0, created_at: $now,
                         valid_from: $now, valid_to: $valid_to
                       })
                       CREATE (d)-[:CONTAINS]->(c)
                       RETURN elementId(c) AS eid""",
                    doc_id=doc_id, cid=cid, text=text, pos=i,
                    title=title, agent_id=agent_id, access_level=access_level,
                    now=now, valid_to=valid_to,
                )
                chunk_eids.append((await c_rec.single())["eid"])

            # Link consecutive chunks with :NEXT for traversal
            for j in range(len(chunk_eids) - 1):
                await s.run(
                    """MATCH (a:Chunk) WHERE elementId(a) = $a
                       MATCH (b:Chunk) WHERE elementId(b) = $b
                       CREATE (a)-[:NEXT]->(b)""",
                    a=chunk_eids[j], b=chunk_eids[j + 1],
                )
        return doc_eid, chunk_eids
    finally:
        await driver.close()


async def _qdrant_ingest(collection: str, chunks: list[str], chunk_eids: list[str],
                         doc_id: str, title: str, agent_id: str, access_level: str,
                         valid_to: str | None = None):
    now = _now_iso()

    # Embed in batches with retry + inter-batch delay
    all_vectors: list[list[float]] = []
    for i in range(0, len(chunks), EMBED_BATCH_SIZE):
        batch = chunks[i:i + EMBED_BATCH_SIZE]
        vecs = await _embed_batch(batch)
        all_vectors.extend(vecs)
        if i + EMBED_BATCH_SIZE < len(chunks):
            await asyncio.sleep(EMBED_BATCH_DELAY)

    # Ensure collection exists
    async with httpx.AsyncClient(timeout=30) as c:
        chk = await c.get(f"{QDRANT_URL}/collections/{collection}")
        if chk.status_code == 404:
            await c.put(
                f"{QDRANT_URL}/collections/{collection}",
                json={"vectors": {"size": len(all_vectors[0]), "distance": "Cosine"}},
            )

    # Upsert points
    points = []
    for i, (text, eid, vec) in enumerate(zip(chunks, chunk_eids, all_vectors)):
        point_id = str(uuid.uuid5(uuid.NAMESPACE_URL, eid))
        points.append({
            "id":      point_id,
            "vector":  vec,
            "payload": {
                "text":             text,
                "neo4j_element_id": eid,
                "node_type":        "Chunk",
                "document_id":      doc_id,
                "document_title":   title,
                "position":         i,
                "agent_id":         agent_id,
                "access_level":     access_level,
                "confidence":       1.0,
                "valid_from":       now,
                "valid_to":         valid_to,
            },
        })

    async with httpx.AsyncClient(timeout=60) as c:
        resp = await c.put(
            f"{QDRANT_URL}/collections/{collection}/points",
            params={"wait": "true"},
            json={"points": points},
        )
        resp.raise_for_status()

    return len(points)


@app.get("/ingest", response_class=HTMLResponse)
async def ingest_form():
    return _UPLOAD_FORM


@app.post("/ingest")
async def ingest_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    title: str = Form(""),
    agent_id: str = Form(""),
    access_level: str = Form("public"),
    valid_hours: float = Form(0.0),
):
    filename = file.filename or "upload"
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Empty file")

    ext = _resolve_ext(filename, data)
    if ext not in _INGEST_ACCEPT:
        raise HTTPException(
            status_code=415,
            detail=f"Unsupported file type '{ext}'. Accepted: {', '.join(sorted(_INGEST_ACCEPT))}",
        )

    doc_title   = title.strip() or filename
    doc_id      = hashlib.md5(f"{filename}-{datetime.now().isoformat()}".encode()).hexdigest()[:16]
    image_props: dict = {}

    try:
        if _is_image(ext):
            text, image_props = await _extract_image_text(filename, data, ext)
        else:
            text = _extract_text(filename, data)
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Text extraction failed: {exc}")

    text = text.strip()
    if not text:
        raise HTTPException(status_code=422, detail="No text could be extracted from the file")

    chunks = _chunk_text(text)
    if not chunks:
        raise HTTPException(status_code=422, detail="File produced no chunks after extraction")

    from datetime import timedelta
    valid_to = (
        (datetime.now(timezone.utc) + timedelta(hours=valid_hours)).isoformat()
        if valid_hours > 0 else None
    )
    logger.info(
        "Ingest '%s': %d chars → %d chunks (ext=%s, valid_to=%s)",
        doc_title, len(text), len(chunks), ext, valid_to,
    )

    try:
        doc_eid, chunk_eids = await _neo4j_ingest(
            doc_id, doc_title, filename, chunks, agent_id, access_level, image_props,
            valid_to=valid_to,
        )
    except Exception as exc:
        logger.error("Neo4j ingest failed: %s", exc)
        raise HTTPException(status_code=502, detail=f"Neo4j ingest failed: {exc}")

    try:
        upserted = await _qdrant_ingest(
            QDRANT_COLLECTION, chunks, chunk_eids, doc_id, doc_title, agent_id, access_level,
            valid_to=valid_to,
        )
    except Exception as exc:
        logger.error("Qdrant ingest failed: %s", exc)
        raise HTTPException(status_code=502, detail=f"Qdrant ingest failed: {exc}")

    background_tasks.add_task(_neo4j_create_mentions, chunk_eids, chunks)

    result = {
        "status":        "ok",
        "document_id":   doc_id,
        "document_eid":  doc_eid,
        "title":         doc_title,
        "filename":      filename,
        "chunks":        len(chunks),
        "qdrant_points": upserted,
        "collection":    QDRANT_COLLECTION,
    }
    if valid_to:
        result["valid_to"] = valid_to
    logger.info("Ingest complete: %s", result)

    return HTMLResponse(_RESULT_TMPL.format(
        color="#4ade80",
        heading=f"✓ Ingested \"{doc_title}\"",
        body=json.dumps(result, indent=2),
    ))

